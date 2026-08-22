"""
Dashboard de Configuración — Exit Poll Venezuela
FastAPI + Jinja2 + Bootstrap 5
Uso: uvicorn app:app --reload
"""

import csv
import os
import random
import re
import sys
import json
import asyncio
import sqlite3
import shutil
import subprocess
import tempfile
import uuid
import io
import traceback
import difflib
import time
import unicodedata
from datetime import datetime, timedelta
from pathlib import Path
from typing import Any

from fastapi import FastAPI, Request, Form, UploadFile, File, HTTPException
from fastapi.responses import HTMLResponse, RedirectResponse, JSONResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates

BASE_DIR = Path(__file__).resolve().parent
DB_PATH = BASE_DIR / "exitpoll.db"
UPLOAD_DIR = BASE_DIR / "static" / "uploads"
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)

app = FastAPI(title="Exit Poll — Configuración")
app.mount("/static", StaticFiles(directory=BASE_DIR / "static"), name="static")
templates = Jinja2Templates(directory=BASE_DIR / "templates")


def _detect_app_version() -> str:
    for key in ("APP_COMMIT_SHA", "GITHUB_SHA", "COMMIT_SHA", "WEBSITE_DEPLOYMENT_ID"):
        value = (os.environ.get(key) or "").strip()
        if value:
            return value[:12]
    try:
        return subprocess.check_output(
            ["git", "rev-parse", "--short=12", "HEAD"],
            cwd=BASE_DIR.parent,
            text=True,
            stderr=subprocess.DEVNULL,
            timeout=2,
        ).strip()
    except Exception:
        return "dev-local"


APP_VERSION = _detect_app_version()
templates.env.globals["app_version"] = APP_VERSION


def _seed_historicos() -> None:
    sys.path.insert(0, str(BASE_DIR))
    import seed_resultados_historicos
    import seed_historico_estudios

    seed_resultados_historicos.seed_resultados_historicos(DB_PATH)
    seed_historico_estudios.seed_historico_estudios(DB_PATH)


async def _seed_historicos_background() -> None:
    try:
        await asyncio.to_thread(_seed_historicos)
        print("[startup] Estudios historicos sembrados en segundo plano")
    except Exception as exc:
        print(f"[startup] WARN: no se pudieron sembrar estudios historicos: {exc}")


@app.on_event("startup")
async def seed_historicos_startup() -> None:
    current = getattr(app.state, "historicos_seed_task", None)
    if current is None or current.done():
        app.state.historicos_seed_task = asyncio.create_task(_seed_historicos_background())


@app.get("/version", response_class=JSONResponse)
async def version_info():
    return {
        "version": APP_VERSION,
        "db_path": str(DB_PATH),
    }


# ── helpers ──────────────────────────────────────────────────────
def get_db() -> sqlite3.Connection:
    conn = sqlite3.connect(str(DB_PATH))
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA foreign_keys = ON")
    return conn


AI_PROVIDER_DEFAULTS = {
    "openai": {"label": "OpenAI", "model": "gpt-4o"},
    "groq": {"label": "Groq", "model": "llama-3.1-8b-instant"},
    "anthropic": {"label": "Anthropic", "model": "claude-sonnet-4-5"},
    "gemini": {"label": "Gemini", "model": "gemini-1.5-pro"},
}

TM_AI_CHUNK_SIZE = 15000
TM_AI_MAX_ATTEMPTS = 5
TM_AI_RETRY_BASE_SECONDS = 2


def ensure_config_table(db: sqlite3.Connection) -> None:
    db.execute("""
        CREATE TABLE IF NOT EXISTS config (
            provider      TEXT PRIMARY KEY,
            api_key       TEXT,
            model         TEXT NOT NULL,
            temperature   REAL NOT NULL DEFAULT 0.3,
            max_tokens    INTEGER NOT NULL DEFAULT 300,
            active        INTEGER NOT NULL DEFAULT 0,
            updated_at    TEXT DEFAULT (datetime('now')),
            CHECK(provider IN ('openai','groq','anthropic','gemini')),
            CHECK(active IN (0,1))
        )
    """)
    active_count = db.execute("SELECT COUNT(*) c FROM config WHERE active=1").fetchone()["c"]
    for provider, defaults in AI_PROVIDER_DEFAULTS.items():
        db.execute("""
            INSERT OR IGNORE INTO config (provider, model, temperature, max_tokens, active)
            VALUES (?, ?, 0, 300, ?)
        """, (provider, defaults["model"], 1 if provider == "openai" and active_count == 0 else 0))
    db.commit()


def ensure_tm_ai_tables(db: sqlite3.Connection) -> None:
    db.execute("""
        CREATE TABLE IF NOT EXISTS election_centers (
            eleccion_id     INTEGER NOT NULL REFERENCES elecciones(id),
            centro_id       TEXT NOT NULL REFERENCES centros(codigo_cne),
            eligible        INTEGER NOT NULL DEFAULT 1,
            source_file     TEXT,
            campos_extra    TEXT,
            created_at      TEXT DEFAULT (datetime('now')),
            updated_at      TEXT DEFAULT (datetime('now')),
            PRIMARY KEY(eleccion_id, centro_id),
            CHECK(eligible IN (0,1))
        )
    """)
    db.execute("CREATE INDEX IF NOT EXISTS idx_ec_eleccion ON election_centers(eleccion_id)")
    db.execute("CREATE INDEX IF NOT EXISTS idx_ec_centro ON election_centers(centro_id)")
    db.execute("""
        CREATE TABLE IF NOT EXISTS tm_ingestion_logs (
            id                  INTEGER PRIMARY KEY,
            eleccion_id          INTEGER NOT NULL REFERENCES elecciones(id),
            source_files         TEXT NOT NULL,
            detected_columns     TEXT,
            field_notes          TEXT,
            match_stats          TEXT,
            user                 TEXT,
            created_at           TEXT DEFAULT (datetime('now'))
        )
    """)
    db.commit()


def get_ai_config(db: sqlite3.Connection, provider: str | None = None) -> sqlite3.Row:
    ensure_config_table(db)
    if provider:
        row = db.execute("SELECT * FROM config WHERE provider=?", (provider,)).fetchone()
    else:
        row = db.execute("SELECT * FROM config WHERE active=1 LIMIT 1").fetchone()
    if not row:
        row = db.execute("SELECT * FROM config WHERE provider='openai'").fetchone()
    return row


def flash(request: Request, msg: str, cat: str = "success"):
    """Almacena un flash message en query-string (stateless)."""
    pass  # usaremos query params ?msg=...&cat=...


# ── INDEX ────────────────────────────────────────────────────────
@app.get("/", response_class=HTMLResponse)
async def index(request: Request):
    db = get_db()
    eleccion = db.execute(
        "SELECT * FROM elecciones WHERE activa = 1 LIMIT 1"
    ).fetchone()
    stats = {}
    if eleccion:
        eid = eleccion["id"]
        stats["candidatos"] = db.execute(
            "SELECT COUNT(*) c FROM candidatos WHERE id_eleccion=?", (eid,)
        ).fetchone()["c"]
        stats["centros_muestra"] = db.execute(
            "SELECT COUNT(*) c FROM muestra WHERE id_eleccion=?", (eid,)
        ).fetchone()["c"]
        stats["centros_con_peso"] = db.execute(
            """SELECT COUNT(*) c FROM pesos p
               JOIN muestra m ON p.id_muestra=m.id
               WHERE m.id_eleccion=?""", (eid,)
        ).fetchone()["c"]
    db.close()
    return templates.TemplateResponse(request=request, name="index.html", context={
        "eleccion": eleccion, "stats": stats
    })


# ══════════════════════════════════════════════════════════════════
# ELECCIONES
# ══════════════════════════════════════════════════════════════════

@app.get("/elecciones", response_class=HTMLResponse)
async def elecciones_list(request: Request, msg: str = "", cat: str = "success"):
    db = get_db()
    rows = db.execute("SELECT * FROM elecciones ORDER BY fecha DESC").fetchall()
    db.close()
    return templates.TemplateResponse(request=request, name="elecciones.html", context={
        "elecciones": rows, "msg": msg, "cat": cat
    })


@app.get("/elecciones/nueva", response_class=HTMLResponse)
async def eleccion_form(request: Request):
    return templates.TemplateResponse(request=request, name="eleccion_form.html", context={
        "eleccion": None
    })


@app.get("/elecciones/{eid}/editar", response_class=HTMLResponse)
async def eleccion_edit(request: Request, eid: int):
    db = get_db()
    row = db.execute("SELECT * FROM elecciones WHERE id=?", (eid,)).fetchone()
    db.close()
    if not row:
        raise HTTPException(404)
    return templates.TemplateResponse(request=request, name="eleccion_form.html", context={
        "eleccion": row
    })


@app.post("/elecciones/guardar")
async def eleccion_save(
    request: Request,
    eid: int = Form(0),
    nombre: str = Form(...),
    tipo: str = Form(...),
    fecha: str = Form(...),
    hora_apertura: str = Form("07:00"),
    hora_cierre: str = Form("18:00"),
    activa: int = Form(0),
):
    db = get_db()
    if activa:
        db.execute("UPDATE elecciones SET activa=0")
    if eid:
        db.execute(
            """UPDATE elecciones
               SET nombre=?, tipo=?, fecha=?, hora_apertura=?, hora_cierre=?, activa=?
               WHERE id=?""",
            (nombre, tipo, fecha, hora_apertura, hora_cierre, activa, eid),
        )
    else:
        db.execute(
            """INSERT INTO elecciones (nombre, tipo, fecha, hora_apertura, hora_cierre, activa)
               VALUES (?,?,?,?,?,?)""",
            (nombre, tipo, fecha, hora_apertura, hora_cierre, activa),
        )
    db.commit()
    db.close()
    return RedirectResponse("/elecciones?msg=Elección+guardada", status_code=303)


@app.post("/elecciones/{eid}/eliminar")
async def eleccion_delete(eid: int):
    db = get_db()
    db.execute("DELETE FROM elecciones WHERE id=?", (eid,))
    db.commit()
    db.close()
    return RedirectResponse("/elecciones?msg=Elección+eliminada&cat=warning", status_code=303)


@app.post("/elecciones/{eid}/activar")
async def eleccion_activar(eid: int):
    db = get_db()
    db.execute("UPDATE elecciones SET activa=0")
    db.execute("UPDATE elecciones SET activa=1 WHERE id=?", (eid,))
    db.commit()
    db.close()
    return RedirectResponse("/elecciones?msg=Elección+activada", status_code=303)


# ══════════════════════════════════════════════════════════════════
# CANDIDATOS
# ══════════════════════════════════════════════════════════════════

@app.get("/candidatos", response_class=HTMLResponse)
async def candidatos_list(request: Request, msg: str = "", cat: str = "success"):
    db = get_db()
    eleccion = db.execute("SELECT * FROM elecciones WHERE activa=1 LIMIT 1").fetchone()
    rows = []
    if eleccion:
        rows = db.execute(
            """SELECT c.*, e.nombre as eleccion_nombre
               FROM candidatos c JOIN elecciones e ON c.id_eleccion=e.id
               WHERE c.id_eleccion=? ORDER BY c.orden""",
            (eleccion["id"],)
        ).fetchall()
    db.close()
    return templates.TemplateResponse(request=request, name="candidatos.html", context={
        "candidatos": rows, "eleccion": eleccion,
        "msg": msg, "cat": cat
    })


@app.get("/candidatos/nuevo", response_class=HTMLResponse)
async def candidato_form(request: Request):
    db = get_db()
    eleccion = db.execute("SELECT * FROM elecciones WHERE activa=1 LIMIT 1").fetchone()
    estados = db.execute("SELECT * FROM estados ORDER BY nombre").fetchall()
    municipios = db.execute(
        """SELECT mu.*, e.nombre AS estado_nombre
           FROM municipios mu
           JOIN estados e ON e.id = mu.id_estado
           ORDER BY e.nombre, mu.nombre"""
    ).fetchall()
    circuitos = db.execute(
        """SELECT ci.*, e.nombre AS estado_nombre
           FROM circuitos ci
           JOIN estados e ON e.id = ci.id_estado
           ORDER BY e.nombre, ci.numero"""
    ).fetchall()
    circ_indigenas = db.execute("SELECT * FROM circunscripciones_indigenas ORDER BY nombre").fetchall()
    db.close()
    if not eleccion:
        return RedirectResponse("/elecciones?msg=Primero+active+una+elección&cat=warning", status_code=303)
    return templates.TemplateResponse(request=request, name="candidato_form.html", context={
        "candidato": None, "eleccion": eleccion,
        "estados": estados,
        "municipios": municipios,
        "circuitos": circuitos,
        "circ_indigenas": circ_indigenas,
    })


@app.get("/candidatos/{cid}/editar", response_class=HTMLResponse)
async def candidato_edit(request: Request, cid: int):
    db = get_db()
    row = db.execute("SELECT * FROM candidatos WHERE id=?", (cid,)).fetchone()
    if not row:
        db.close()
        raise HTTPException(404)
    eleccion = db.execute("SELECT * FROM elecciones WHERE id=?", (row["id_eleccion"],)).fetchone()
    estados = db.execute("SELECT * FROM estados ORDER BY nombre").fetchall()
    municipios = db.execute(
        """SELECT mu.*, e.nombre AS estado_nombre
           FROM municipios mu
           JOIN estados e ON e.id = mu.id_estado
           ORDER BY e.nombre, mu.nombre"""
    ).fetchall()
    circuitos = db.execute(
        """SELECT ci.*, e.nombre AS estado_nombre
           FROM circuitos ci
           JOIN estados e ON e.id = ci.id_estado
           ORDER BY e.nombre, ci.numero"""
    ).fetchall()
    circ_indigenas = db.execute("SELECT * FROM circunscripciones_indigenas ORDER BY nombre").fetchall()
    db.close()
    return templates.TemplateResponse(request=request, name="candidato_form.html", context={
        "candidato": row, "eleccion": eleccion,
        "estados": estados,
        "municipios": municipios,
        "circuitos": circuitos,
        "circ_indigenas": circ_indigenas,
    })


@app.post("/candidatos/guardar")
async def candidato_save(
    request: Request,
    cid: int = Form(0),
    id_eleccion: int = Form(...),
    nombre: str = Form(...),
    partido: str = Form(""),
    bando: str = Form("otro"),
    tipo: str = Form("unico"),
    id_estado: int = Form(0),
    id_municipio: int = Form(0),
    id_circuito: int = Form(0),
    id_circ_indigena: int = Form(0),
    orden: int = Form(1),
    foto: UploadFile = File(None),
):
    id_estado_db = id_estado or None
    id_municipio_db = id_municipio or None
    id_circuito_db = id_circuito or None
    id_circ_indigena_db = id_circ_indigena or None
    db_meta = get_db()
    try:
        eleccion_meta = db_meta.execute(
            "SELECT tipo FROM elecciones WHERE id=?", (id_eleccion,)
        ).fetchone()
        tipo_eleccion = eleccion_meta["tipo"] if eleccion_meta else "nacional"
    finally:
        db_meta.close()

    if tipo == "unico" and tipo_eleccion == "regional":
        id_municipio_db = id_circuito_db = id_circ_indigena_db = None
    elif tipo == "unico" and tipo_eleccion == "municipal":
        id_estado_db = id_circuito_db = id_circ_indigena_db = None
    elif tipo == "unico":
        id_estado_db = id_municipio_db = id_circuito_db = id_circ_indigena_db = None
    elif tipo == "lista":
        id_municipio_db = id_circuito_db = id_circ_indigena_db = None
    elif tipo == "nominal":
        id_municipio_db = id_circ_indigena_db = None
    elif tipo == "indigena":
        id_estado_db = id_municipio_db = id_circuito_db = None

    foto_url = None
    if foto and foto.filename:
        ext = Path(foto.filename).suffix.lower()
        if ext not in (".jpg", ".jpeg", ".png", ".webp"):
            return RedirectResponse(
                "/candidatos?msg=Formato+de+imagen+no+válido&cat=danger", status_code=303
            )
        fname = f"{uuid.uuid4().hex}{ext}"
        dest = UPLOAD_DIR / fname
        with open(dest, "wb") as f:
            shutil.copyfileobj(foto.file, f)
        foto_url = f"/static/uploads/{fname}"

    db = get_db()
    if cid:
        if foto_url:
            db.execute(
                """UPDATE candidatos
                   SET nombre=?, partido=?, bando=?, tipo=?,
                       id_estado=?, id_municipio=?, id_circuito=?, id_circ_indigena=?,
                       orden=?, foto_url=?
                   WHERE id=?""",
                (
                    nombre, partido, bando, tipo,
                    id_estado_db, id_municipio_db, id_circuito_db, id_circ_indigena_db,
                    orden, foto_url, cid,
                ),
            )
        else:
            db.execute(
                """UPDATE candidatos
                   SET nombre=?, partido=?, bando=?, tipo=?,
                       id_estado=?, id_municipio=?, id_circuito=?, id_circ_indigena=?,
                       orden=?
                   WHERE id=?""",
                (
                    nombre, partido, bando, tipo,
                    id_estado_db, id_municipio_db, id_circuito_db, id_circ_indigena_db,
                    orden, cid,
                ),
            )
    else:
        db.execute(
            """INSERT INTO candidatos (
                   id_eleccion, nombre, partido, bando, tipo,
                   id_estado, id_municipio, id_circuito, id_circ_indigena,
                   orden, foto_url
               )
               VALUES (?,?,?,?,?,?,?,?,?,?,?)""",
            (
                id_eleccion, nombre, partido, bando, tipo,
                id_estado_db, id_municipio_db, id_circuito_db, id_circ_indigena_db,
                orden, foto_url,
            ),
        )
    db.commit()
    db.close()
    return RedirectResponse("/candidatos?msg=Candidato+guardado", status_code=303)


@app.post("/candidatos/{cid}/eliminar")
async def candidato_delete(cid: int):
    db = get_db()
    db.execute("DELETE FROM candidatos WHERE id=?", (cid,))
    db.commit()
    db.close()
    return RedirectResponse("/candidatos?msg=Candidato+eliminado&cat=warning", status_code=303)


# ══════════════════════════════════════════════════════════════════
# FICHA TÉCNICA DE LA MUESTRA
# ══════════════════════════════════════════════════════════════════

@app.get("/ficha", response_class=HTMLResponse)
async def ficha_tecnica(request: Request):
    db = get_db()
    eleccion = db.execute("SELECT * FROM elecciones WHERE activa=1 LIMIT 1").fetchone()

    # --- Registro Electoral (siempre, independiente de elección/muestra) ---
    re = {}
    re["centros_activos"] = db.execute(
        "SELECT COUNT(*) c FROM centros WHERE activo=1"
    ).fetchone()["c"]
    re["centros_inactivos"] = db.execute(
        "SELECT COUNT(*) c FROM centros WHERE activo=0"
    ).fetchone()["c"]
    re["electores"] = db.execute(
        "SELECT COALESCE(SUM(num_electores),0) s FROM centros WHERE activo=1"
    ).fetchone()["s"]
    re["mesas"] = db.execute(
        "SELECT COALESCE(SUM(num_mesas),0) s FROM centros WHERE activo=1"
    ).fetchone()["s"]
    re["con_gps"] = db.execute(
        "SELECT COUNT(*) c FROM centros WHERE lat IS NOT NULL AND lon IS NOT NULL AND activo=1"
    ).fetchone()["c"]
    re["total_estados"] = db.execute("SELECT COUNT(*) c FROM estados").fetchone()["c"]
    re["estados_activos"] = db.execute(
        "SELECT COUNT(DISTINCT id_estado) c FROM centros WHERE activo=1"
    ).fetchone()["c"]

    # Desglose RE por estado (incluye estados con centros inactivos)
    re_estados = db.execute(
        """SELECT e.nombre as estado,
                  SUM(CASE WHEN c.activo=1 THEN 1 ELSE 0 END) as centros_activos,
                  SUM(CASE WHEN c.activo=0 THEN 1 ELSE 0 END) as centros_inactivos,
                  SUM(CASE WHEN c.activo=1 THEN c.num_electores ELSE 0 END) as electores,
                  SUM(CASE WHEN c.activo=1 THEN c.num_mesas ELSE 0 END) as mesas,
                  SUM(CASE WHEN c.lat IS NOT NULL AND c.activo=1 THEN 1 ELSE 0 END) as con_gps
           FROM estados e
           LEFT JOIN centros c ON c.id_estado=e.id
           GROUP BY e.id ORDER BY e.nombre"""
    ).fetchall()

    # --- Muestra (solo si hay elección activa con muestra) ---
    muestra = {}
    muestra_estados = []
    if eleccion:
        eid = eleccion["id"]
        muestra["total_centros"] = db.execute(
            "SELECT COUNT(*) c FROM muestra WHERE id_eleccion=? AND activo=1", (eid,)
        ).fetchone()["c"]
        if muestra["total_centros"] > 0:
            muestra["electores"] = db.execute(
                """SELECT COALESCE(SUM(ct.num_electores),0) s
                   FROM muestra m JOIN centros ct ON m.codigo_centro=ct.codigo_cne
                   WHERE m.id_eleccion=? AND m.activo=1""", (eid,)
            ).fetchone()["s"]
            muestra["mesas"] = db.execute(
                """SELECT COALESCE(SUM(ct.num_mesas),0) s
                   FROM muestra m JOIN centros ct ON m.codigo_centro=ct.codigo_cne
                   WHERE m.id_eleccion=? AND m.activo=1""", (eid,)
            ).fetchone()["s"]
            if re["electores"]:
                muestra["cobertura_pct"] = round(
                    100 * muestra["electores"] / re["electores"], 2)
            else:
                muestra["cobertura_pct"] = 0
            muestra["estados_cubiertos"] = db.execute(
                """SELECT COUNT(DISTINCT ct.id_estado) c
                   FROM muestra m JOIN centros ct ON m.codigo_centro=ct.codigo_cne
                   WHERE m.id_eleccion=? AND m.activo=1""", (eid,)
            ).fetchone()["c"]
            muestra_estados = db.execute(
                """SELECT e.nombre as estado, COUNT(m.id) as centros,
                          SUM(ct.num_electores) as electores, SUM(ct.num_mesas) as mesas
                   FROM muestra m
                   JOIN centros ct ON m.codigo_centro=ct.codigo_cne
                   JOIN estados e ON ct.id_estado=e.id
                   WHERE m.id_eleccion=? AND m.activo=1
                   GROUP BY e.id ORDER BY e.nombre""", (eid,)
            ).fetchall()
            muestra["tipos"] = db.execute(
                """SELECT tipo_centro, COUNT(*) c
                   FROM muestra WHERE id_eleccion=? AND activo=1 AND tipo_centro IS NOT NULL
                   GROUP BY tipo_centro""", (eid,)
            ).fetchall()

    # --- Error muestral ---
    # Formula: e = Z * sqrt(P*Q/n) * sqrt(1 - n/N)
    # Z=1.96 (95%), P=Q=0.5
    error_muestral = None
    if muestra.get("total_centros", 0) > 0 and muestra.get("electores", 0) > 0 and re["electores"] > 0:
        import math
        n = muestra["electores"]
        N = re["electores"]
        error_muestral = 1.96 * math.sqrt(0.25 / n) * math.sqrt(1 - n / N) * 100

    db.close()
    return templates.TemplateResponse(request=request, name="ficha.html", context={
        "eleccion": eleccion,
        "re": re, "re_estados": re_estados,
        "muestra": muestra, "muestra_estados": muestra_estados,
        "error_muestral": error_muestral
    })


# ══════════════════════════════════════════════════════════════════
# PESOS POR CENTRO
# ══════════════════════════════════════════════════════════════════

@app.get("/pesos", response_class=HTMLResponse)
async def pesos_list(request: Request, estado: str = "", msg: str = "", cat: str = "success"):
    db = get_db()
    eleccion = db.execute("SELECT * FROM elecciones WHERE activa=1 LIMIT 1").fetchone()
    rows = []
    estados = []
    tiene_muestra = False

    if eleccion:
        eid = eleccion["id"]
        tiene_muestra = db.execute(
            "SELECT COUNT(*) c FROM muestra WHERE id_eleccion=?", (eid,)
        ).fetchone()["c"] > 0

    if eleccion and tiene_muestra:
        # Estados con centros en la muestra
        estados = db.execute(
            """SELECT DISTINCT e.id, e.nombre FROM estados e
               JOIN centros ct ON ct.id_estado=e.id
               JOIN muestra m ON m.codigo_centro=ct.codigo_cne
               WHERE m.id_eleccion=?
               ORDER BY e.nombre""", (eid,)
        ).fetchall()

        # Centros de la muestra con sus pesos
        query = """
            SELECT m.id as id_muestra, m.codigo_centro, ct.nombre as centro_nombre,
                   e.nombre as estado, mu.nombre as municipio, p2.nombre as parroquia,
                   m.tipo_centro, m.activo,
                   COALESCE(p.peso_parroquia,0) as peso_parroquia,
                   COALESCE(p.peso_municipio,0) as peso_municipio,
                   COALESCE(p.peso_estado,0) as peso_estado,
                   COALESCE(p.peso_nacion,0) as peso_nacion,
                   ct.num_electores, ct.num_mesas
            FROM muestra m
            JOIN centros ct ON m.codigo_centro=ct.codigo_cne
            JOIN estados e ON ct.id_estado=e.id
            LEFT JOIN municipios mu ON ct.id_municipio=mu.id
            LEFT JOIN parroquias p2 ON ct.id_parroquia=p2.id
            LEFT JOIN pesos p ON p.id_muestra=m.id
            WHERE m.id_eleccion=? AND m.activo=1
        """
        params = [eid]
        if estado:
            query += " AND e.id=?"
            params.append(int(estado))
        query += " ORDER BY e.nombre, mu.nombre, ct.nombre"
        rows = db.execute(query, params).fetchall()

    db.close()
    return templates.TemplateResponse(request=request, name="pesos.html", context={
        "eleccion": eleccion, "pesos": rows,
        "estados": estados, "estado_sel": estado, "msg": msg, "cat": cat,
        "tiene_muestra": tiene_muestra
    })


@app.post("/pesos/{id_muestra}/guardar")
async def peso_save(
    id_muestra: int,
    peso_parroquia: float = Form(0),
    peso_municipio: float = Form(0),
    peso_estado: float = Form(0),
    peso_nacion: float = Form(0),
):
    db = get_db()
    exists = db.execute("SELECT 1 FROM pesos WHERE id_muestra=?", (id_muestra,)).fetchone()
    if exists:
        db.execute(
            """UPDATE pesos SET peso_parroquia=?, peso_municipio=?, peso_estado=?, peso_nacion=?
               WHERE id_muestra=?""",
            (peso_parroquia, peso_municipio, peso_estado, peso_nacion, id_muestra),
        )
    else:
        db.execute(
            """INSERT INTO pesos (id_muestra, peso_parroquia, peso_municipio, peso_estado, peso_nacion)
               VALUES (?,?,?,?,?)""",
            (id_muestra, peso_parroquia, peso_municipio, peso_estado, peso_nacion),
        )
    db.commit()
    db.close()
    return RedirectResponse("/pesos?msg=Peso+actualizado", status_code=303)


@app.post("/pesos/guardar-todos")
async def pesos_save_all(request: Request):
    """Guarda todos los pesos editados desde la tabla inline."""
    form = await request.form()
    db = get_db()
    count = 0
    for key, val in form.items():
        if key.startswith("pp_"):
            mid = int(key[3:])
            pp = float(val or 0)
            pm = float(form.get(f"pm_{mid}", 0) or 0)
            pe = float(form.get(f"pe_{mid}", 0) or 0)
            pn = float(form.get(f"pn_{mid}", 0) or 0)
            exists = db.execute("SELECT 1 FROM pesos WHERE id_muestra=?", (mid,)).fetchone()
            if exists:
                db.execute(
                    """UPDATE pesos SET peso_parroquia=?, peso_municipio=?, peso_estado=?, peso_nacion=?
                       WHERE id_muestra=?""",
                    (pp, pm, pe, pn, mid),
                )
            else:
                db.execute(
                    """INSERT INTO pesos (id_muestra, peso_parroquia, peso_municipio, peso_estado, peso_nacion)
                       VALUES (?,?,?,?,?)""",
                    (mid, pp, pm, pe, pn),
                )
            count += 1
    db.commit()
    db.close()
    return RedirectResponse(f"/pesos?msg=Pesos+actualizados+({count}+centros)&cat=success", status_code=303)


@app.get("/pesos/{id_muestra}/editar", response_class=HTMLResponse)
async def peso_edit_form(request: Request, id_muestra: int):
    db = get_db()
    row = db.execute(
        """SELECT m.id as id_muestra, m.codigo_centro, ct.nombre as centro_nombre,
                  e.nombre as estado, mu.nombre as municipio,
                  COALESCE(p.peso_parroquia,0) as peso_parroquia,
                  COALESCE(p.peso_municipio,0) as peso_municipio,
                  COALESCE(p.peso_estado,0) as peso_estado,
                  COALESCE(p.peso_nacion,0) as peso_nacion,
                  ct.num_electores
           FROM muestra m
           JOIN centros ct ON m.codigo_centro=ct.codigo_cne
           JOIN estados e ON ct.id_estado=e.id
           LEFT JOIN municipios mu ON ct.id_municipio=mu.id
           LEFT JOIN pesos p ON p.id_muestra=m.id
           WHERE m.id=?""", (id_muestra,)
    ).fetchone()
    db.close()
    if not row:
        raise HTTPException(404)
    return templates.TemplateResponse(request=request, name="peso_edit.html", context={
        "centro": row
    })


# ══════════════════════════════════════════════════════════════════
# TABLA DE MESA — Carga y gestión
# ══════════════════════════════════════════════════════════════════

@app.get("/tm", response_class=HTMLResponse)
async def tm_index(request: Request, msg: str = "", cat: str = "success"):
    db = get_db()
    ensure_tm_ai_tables(db)
    stats = {}
    stats["centros_activos"] = db.execute(
        "SELECT COUNT(*) c FROM centros WHERE activo=1"
    ).fetchone()["c"]
    stats["centros_inactivos"] = db.execute(
        "SELECT COUNT(*) c FROM centros WHERE activo=0"
    ).fetchone()["c"]
    stats["con_gps"] = db.execute(
        "SELECT COUNT(*) c FROM centros WHERE lat IS NOT NULL AND lon IS NOT NULL"
    ).fetchone()["c"]
    stats["electores"] = db.execute(
        "SELECT COALESCE(SUM(num_electores),0) s FROM centros WHERE activo=1"
    ).fetchone()["s"]
    stats["estados"] = db.execute("SELECT COUNT(*) c FROM estados").fetchone()["c"]
    stats["municipios"] = db.execute("SELECT COUNT(*) c FROM municipios").fetchone()["c"]
    stats["parroquias"] = db.execute("SELECT COUNT(*) c FROM parroquias").fetchone()["c"]
    elecciones = db.execute("SELECT id, nombre, tipo, fecha, activa FROM elecciones ORDER BY activa DESC, fecha DESC").fetchall()

    # Resumen por estado
    por_estado_raw = db.execute(
        """SELECT e.nombre, COUNT(c.codigo_cne) as centros,
                  SUM(c.num_electores) as electores, SUM(c.num_mesas) as mesas,
                  SUM(CASE WHEN c.lat IS NOT NULL THEN 1 ELSE 0 END) as con_gps
           FROM centros c
           JOIN estados e ON c.id_estado=e.id
           WHERE c.activo=1
           GROUP BY e.id ORDER BY e.nombre"""
    ).fetchall()
    por_estado_map = {}
    for row in por_estado_raw:
        key = _canonical_estado_name(row["nombre"])
        nombre = "LA GUAIRA" if key == "LA GUAIRA" else row["nombre"]
        agg = por_estado_map.setdefault(nombre, {"nombre": nombre, "centros": 0, "electores": 0, "mesas": 0, "con_gps": 0})
        agg["centros"] += int(row["centros"] or 0)
        agg["electores"] += int(row["electores"] or 0)
        agg["mesas"] += int(row["mesas"] or 0)
        agg["con_gps"] += int(row["con_gps"] or 0)
    por_estado = sorted(por_estado_map.values(), key=lambda r: r["nombre"])
    db.close()
    return templates.TemplateResponse(request=request, name="tm.html", context={
        "stats": stats, "por_estado": por_estado,
        "elecciones": elecciones,
        "msg": msg, "cat": cat
    })


@app.post("/tm/cargar")
async def tm_upload(
    request: Request,
    archivo: UploadFile = File(...),
    formato: str = Form("auto"),
    hoja: str = Form(""),
    dry_run: int = Form(0),
):
    """Recibe un archivo TM (Excel o CSV), lo convierte y carga en la BD."""
    # Guardar archivo temporal
    ext = Path(archivo.filename).suffix.lower()
    if ext not in (".xlsx", ".xls", ".xlsm", ".csv"):
        return RedirectResponse(
            "/tm?msg=Formato+no+soportado.+Use+Excel+o+CSV&cat=danger",
            status_code=303,
        )

    tmp_path = UPLOAD_DIR / f"tm_upload_{uuid.uuid4().hex}{ext}"
    with open(tmp_path, "wb") as f:
        shutil.copyfileobj(archivo.file, f)

    output_lines = []
    try:
        # Importar convertidor y cargador
        sys.path.insert(0, str(BASE_DIR))
        import convertidor_tm
        import cargador_tm

        # Paso 1: si es Excel, convertir a CSV estándar
        if ext in (".xlsx", ".xls", ".xlsm"):
            csv_path = tmp_path.with_suffix(".csv")
            df_raw = convertidor_tm.cargar_archivo(str(tmp_path), hoja=hoja or None)
            fmt = formato
            if fmt == "auto":
                fmt = convertidor_tm.detectar_formato(df_raw)
            output_lines.append(f"Formato detectado: {fmt}")

            df = convertidor_tm.CONVERTIDORES[fmt](df_raw)
            df = convertidor_tm.limpiar(df)

            centros = df["codigo_centro"].nunique()
            electores = df["electores"].sum()
            output_lines.append(f"Convertido: {len(df):,} mesas, {centros:,} centros, {electores:,} electores")

            df.to_csv(str(csv_path), index=False, encoding="utf-8-sig")
        else:
            csv_path = tmp_path

        # Paso 2: cargar en BD
        # Capturar output del cargador
        import io as _io
        from contextlib import redirect_stdout
        buf = _io.StringIO()
        with redirect_stdout(buf):
            cargador_tm.cargar_tm(str(csv_path), dry_run=bool(dry_run))
        output_lines.append(buf.getvalue())

        # Limpiar archivos temporales
        tmp_path.unlink(missing_ok=True)
        if csv_path != tmp_path:
            csv_path.unlink(missing_ok=True)

        result_msg = " | ".join(output_lines)
        return RedirectResponse(
            f"/tm?msg={result_msg[:500]}&cat=success",
            status_code=303,
        )

    except Exception as e:
        tmp_path.unlink(missing_ok=True)
        error_msg = str(e)[:300]
        return RedirectResponse(
            f"/tm?msg=Error:+{error_msg}&cat=danger",
            status_code=303,
        )


TM_AI_SYSTEM_PROMPT = """
This is a Venezuelan CNE electoral registry file (tabla de mesa). Column names and available fields vary by election year and event type. Identify all columns present, inspect the actual values to understand what each column represents, map them to the target schema by meaning rather than by name, and extract every data row.

Do not assume a fixed column structure. Infer fields from headers and values. Pay special attention to columns whose values look like center identifiers even if the header says something else. For example, a column named "CODIGO COMUNA CNE" with values like COM_130101001 may be serving as the center identifier for this election; preserve it in campos_extra and note the ambiguity.

Target internal schema keys for each row:
- estado
- cod_estado
- municipio
- cod_municipio
- parroquia
- cod_parroquia
- nombre_centro
- codigo_centro
- num_mesas
- num_electores
- direccion
- campos_extra

Rules:
- Use null for internal schema fields not present.
- Put every source column that does not map to the target schema into campos_extra under its original column name.
- Skip title rows, repeated headers, page numbers, totals, footers, and artifacts.
- Handle encoding artifacts, uppercase text, accents, and abbreviations such as U.E., E.B., MP., PQ.
- Treat VARGAS and LA GUAIRA as the same state name when either appears.
- Return ONLY a JSON object, with no preamble and no markdown fences.
- The JSON object must have exactly these top-level keys: detected_columns, field_notes, centros, match_hints.
"""


def _extract_json_object(text: str) -> dict[str, Any]:
    clean = (text or "").strip()
    clean = re.sub(r"^```(?:json)?\s*", "", clean, flags=re.I)
    clean = re.sub(r"\s*```$", "", clean)
    start = clean.find("{")
    end = clean.rfind("}")
    if start >= 0 and end > start:
        clean = clean[start:end + 1]
    return json.loads(clean)


def _chunk_text(text: str, chunk_size: int = TM_AI_CHUNK_SIZE) -> list[str]:
    text = text or ""
    if len(text) <= chunk_size:
        return [text]
    chunks, pos = [], 0
    while pos < len(text):
        end = min(len(text), pos + chunk_size)
        cut = text.rfind("\n", pos, end)
        if cut <= pos + 1000:
            cut = end
        chunks.append(text[pos:cut])
        pos = cut
    return chunks


def _ai_retry_delay(attempt: int, exc: Exception) -> float:
    retry_after = getattr(exc, "retry_after", None)
    if retry_after is None and getattr(exc, "response", None) is not None:
        headers = getattr(exc.response, "headers", {}) or {}
        retry_after = headers.get("retry-after") or headers.get("Retry-After")
    try:
        delay = float(retry_after) if retry_after is not None else TM_AI_RETRY_BASE_SECONDS ** attempt
    except (TypeError, ValueError):
        delay = TM_AI_RETRY_BASE_SECONDS ** attempt
    return min(delay, 30.0) + random.uniform(0.0, 0.75)


def _read_upload_text(path: Path, ext: str) -> str:
    if ext in (".txt", ".csv"):
        for enc in ("utf-8-sig", "utf-8", "latin-1"):
            try:
                return path.read_text(encoding=enc)
            except UnicodeDecodeError:
                continue
        return path.read_text(errors="ignore")
    if ext in (".xlsx", ".xlsm", ".xls"):
        from openpyxl import load_workbook

        wb = load_workbook(path, read_only=True, data_only=True)
        lines = []
        for ws in wb.worksheets:
            lines.append(f"### SHEET: {ws.title}")
            for row in ws.iter_rows(values_only=True):
                values = ["" if v is None else str(v) for v in row]
                if any(v.strip() for v in values):
                    lines.append("\t".join(values))
        return "\n".join(lines)
    if ext == ".docx":
        try:
            import docx
        except ImportError as exc:
            raise HTTPException(500, "Falta dependencia python-docx para leer DOCX") from exc
        doc = docx.Document(str(path))
        lines = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                lines.append("\t".join(cell.text.strip() for cell in row.cells))
        return "\n".join(lines)
    if ext == ".pdf":
        try:
            import pdfplumber
        except ImportError as exc:
            raise HTTPException(500, "Falta dependencia pdfplumber para leer PDF") from exc
        lines = []
        with pdfplumber.open(str(path)) as pdf:
            for i, page in enumerate(pdf.pages, 1):
                lines.append(f"### PAGE {i}")
                lines.append(page.extract_text() or "")
        return "\n".join(lines)
    raise HTTPException(400, "Formato no soportado")


def _normalize_match_text(value: Any) -> str:
    text = "" if value is None else str(value)
    text = unicodedata.normalize("NFKD", text).encode("ascii", "ignore").decode("ascii")
    text = text.upper()
    replacements = {
        "U.E.": "UNIDAD EDUCATIVA",
        "UE ": "UNIDAD EDUCATIVA ",
        "E.B.": "ESCUELA BASICA",
        "EB ": "ESCUELA BASICA ",
        "MP.": "MUNICIPIO",
        "MCPIO.": "MUNICIPIO",
        "PQ.": "PARROQUIA",
        "PQA.": "PARROQUIA",
        "EDO.": "ESTADO",
        "DTTO.": "DISTRITO",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    text = re.sub(r"[^A-Z0-9]+", " ", text)
    return re.sub(r"\s+", " ", text).strip()


def _center_match_blob(row: sqlite3.Row) -> str:
    return _normalize_match_text(f"{row['nombre']} {row['municipio'] or ''} {row['parroquia'] or ''}")


def _source_match_blob(centro: dict[str, Any]) -> str:
    return _normalize_match_text(
        f"{centro.get('nombre_centro') or ''} {centro.get('municipio') or ''} {centro.get('parroquia') or ''}"
    )


_GEO_PREFIJOS = ("MUNICIPIO ", "PARROQUIA ", "ESTADO ", "DISTRITO ", "CM ", "CIUDAD ")


def _geo_match_name(value: Any) -> str:
    """Nombre canónico de municipio/parroquia para comparar entre fuentes.

    Iguala el nombre crudo del CNE que guarda cargador_tm ("MP. ZAMORA",
    "CM. VILLA DE CURA") con el que extrae la ingesta IA ("Zamora")."""
    name = _normalize_match_text(value)
    for pref in _GEO_PREFIJOS:
        if name.startswith(pref):
            name = name[len(pref):]
            break
    return name.strip()


def _canonical_estado_name(value: Any) -> str:
    name = _normalize_match_text(value or "")
    if name in {"VARGAS", "EDO VARGAS", "ESTADO VARGAS", "LA GUAIRA", "EDO LA GUAIRA", "ESTADO LA GUAIRA"}:
        return "LA GUAIRA"
    return name


def _normalize_center_code(value: Any) -> str:
    raw = str(value or "").strip()
    if not raw:
        return ""
    first_part = raw.split(".", 1)[0].strip()
    if first_part.isdigit() and len(first_part) >= 6:
        return first_part
    digits = re.sub(r"\D+", "", raw)
    return digits or raw.upper()


def _digits_code(value: Any) -> str:
    return re.sub(r"\D+", "", str(value or ""))


def _geo_code_prefix(centro: dict[str, Any]) -> str:
    estado = _digits_code(centro.get("cod_estado"))
    municipio = _digits_code(centro.get("cod_municipio"))
    parroquia = _digits_code(centro.get("cod_parroquia"))
    if estado:
        prefix = estado.zfill(2)[-2:]
        if municipio:
            prefix += municipio.zfill(2)[-2:]
        if parroquia:
            prefix += parroquia.zfill(2)[-2:]
        return prefix

    codigo = _normalize_center_code(centro.get("codigo_centro") or centro.get("codigo_cne"))
    if codigo.isdigit():
        if len(codigo) >= 6:
            return codigo[:6]
        if len(codigo) >= 4:
            return codigo[:4]
        if len(codigo) >= 2:
            return codigo[:2]
    return ""


def _estado_code_from_centro(centro: dict[str, Any], codigo: str | None = None) -> str:
    estado = _digits_code(centro.get("cod_estado"))
    if estado:
        return estado.zfill(2)[-2:]
    codigo_norm = _normalize_center_code(codigo or centro.get("codigo_centro") or centro.get("codigo_cne"))
    if codigo_norm.isdigit() and len(codigo_norm) >= 2:
        return codigo_norm[:2]
    return ""


def _estado_rows(db: sqlite3.Connection) -> list[sqlite3.Row]:
    return db.execute("SELECT id, codigo_cne, nombre FROM estados").fetchall()


def _estado_ids_by_canonical(db: sqlite3.Connection, canonical_name: str) -> set[int]:
    if not canonical_name:
        return set()
    return {
        row["id"] for row in _estado_rows(db)
        if _canonical_estado_name(row["nombre"]) == canonical_name
    }


def _estado_row_for_centro(db: sqlite3.Connection, centro: dict[str, Any], codigo: str | None = None) -> sqlite3.Row | None:
    estado_code = _estado_code_from_centro(centro, codigo)
    if estado_code:
        for row in _estado_rows(db):
            if _digits_code(row["codigo_cne"]).zfill(2)[-2:] == estado_code:
                return row

    canonical_name = _canonical_estado_name(centro.get("estado") or "")
    if canonical_name:
        for row in _estado_rows(db):
            if _canonical_estado_name(row["nombre"]) == canonical_name:
                return row
    return None


def _centro_from_confirm_item(item: dict[str, Any]) -> dict[str, Any]:
    centro = item.get("centro")
    return centro if isinstance(centro, dict) else {}


GENERIC_MATCH_TOKENS = {
    "CENTRO", "ELECTORAL", "UNIDAD", "EDUCATIVA", "ESCUELA", "BASICA",
    "LICEO", "NACIONAL", "BOLIVARIANA", "GRUPO", "COLEGIO", "MUNICIPIO",
    "PARROQUIA", "ESTADO",
}


def _match_tokens(text: str) -> set[str]:
    return {
        token for token in text.split()
        if len(token) >= 3 and token not in GENERIC_MATCH_TOKENS
    }


def _registry_entry(row: sqlite3.Row) -> dict[str, Any]:
    item = dict(row)
    blob = _center_match_blob(row)
    item["_codigo_norm"] = _normalize_center_code(row["codigo_cne"])
    item["_codigo_prefix_2"] = item["_codigo_norm"][:2]
    item["_codigo_prefix_4"] = item["_codigo_norm"][:4]
    item["_codigo_prefix_6"] = item["_codigo_norm"][:6]
    item["_match_blob"] = blob
    item["_tokens"] = _match_tokens(blob)
    item["_estado_norm"] = _canonical_estado_name(row["estado"] or "")
    item["_municipio_norm"] = _normalize_match_text(row["municipio"] or "")
    item["_parroquia_norm"] = _normalize_match_text(row["parroquia"] or "")
    return item


def _candidate_registry(centro: dict[str, Any], registry: list[dict[str, Any]], source_tokens: set[str]) -> list[dict[str, Any]]:
    estado = _canonical_estado_name(centro.get("estado") or "")
    municipio = _normalize_match_text(centro.get("municipio") or "")
    parroquia = _normalize_match_text(centro.get("parroquia") or "")
    code_prefix = _geo_code_prefix(centro)

    pool = registry
    if code_prefix:
        prefix_key = f"_codigo_prefix_{len(code_prefix)}"
        by_code = [r for r in registry if r.get(prefix_key) == code_prefix]
        if by_code:
            pool = by_code

    if municipio:
        by_municipio = [r for r in registry if r["_municipio_norm"] == municipio]
        if by_municipio:
            pool = [r for r in pool if r in by_municipio] or by_municipio
    elif estado:
        by_estado = [r for r in registry if r["_estado_norm"] == estado]
        if by_estado:
            pool = [r for r in pool if r in by_estado] or by_estado

    if parroquia:
        by_parroquia = [r for r in pool if r["_parroquia_norm"] == parroquia]
        if by_parroquia:
            pool = by_parroquia

    if not source_tokens:
        return pool[:500]

    scored = []
    for row in pool:
        overlap = len(source_tokens & row["_tokens"])
        if overlap:
            scored.append((overlap, row))
    if scored:
        scored.sort(key=lambda x: x[0], reverse=True)
        return [row for _, row in scored[:500]]
    return pool[:500]


def _best_fuzzy_match(centro: dict[str, Any], registry: list[dict[str, Any]]) -> tuple[dict[str, Any] | None, float, list[dict]]:
    source = _source_match_blob(centro)
    if not source:
        return None, 0.0, []
    source_tokens = _match_tokens(source)
    candidates_pool = _candidate_registry(centro, registry, source_tokens)
    scored = []
    for row in candidates_pool:
        score = difflib.SequenceMatcher(None, source, row["_match_blob"]).ratio()
        scored.append((score, row))
    scored.sort(key=lambda x: x[0], reverse=True)
    candidates = [
        {
            "codigo_centro": row["codigo_cne"],
            "nombre": row["nombre"],
            "municipio": row["municipio"],
            "parroquia": row["parroquia"],
            "confidence_score": round(score, 3),
        }
        for score, row in scored[:5]
    ]
    return (scored[0][1], round(scored[0][0], 3), candidates) if scored else (None, 0.0, [])


def _registry_rows(db: sqlite3.Connection) -> list[sqlite3.Row]:
    return db.execute("""
        SELECT c.codigo_cne, c.nombre, c.direccion, c.num_mesas, c.num_electores,
               e.nombre AS estado, mu.nombre AS municipio, p.nombre AS parroquia
        FROM centros c
        JOIN estados e ON e.id = c.id_estado
        LEFT JOIN municipios mu ON mu.id = c.id_municipio
        LEFT JOIN parroquias p ON p.id = c.id_parroquia
    """).fetchall()


def _match_centros(db: sqlite3.Connection, centros: list[dict], hints: Any | None = None) -> dict[str, Any]:
    registry = [_registry_entry(row) for row in _registry_rows(db)]
    by_code = {r["_codigo_norm"]: r for r in registry if r["_codigo_norm"]}
    rows = []
    stats = {"MATCHED": 0, "NEW": 0, "AMBIGUOUS": 0, "CONFLICT": 0, "EXTRACTION_ERROR": 0}

    for idx, centro in enumerate(centros):
        if not isinstance(centro, dict) or not (centro.get("nombre_centro") or centro.get("codigo_centro")):
            status = "EXTRACTION_ERROR"
            rows.append({"row_index": idx, "match_status": status, "confidence_score": 0, "centro": centro, "candidates": []})
            stats[status] += 1
            continue

        codigo = _normalize_center_code(centro.get("codigo_centro") or centro.get("codigo_cne"))
        exact = by_code.get(codigo)
        if exact:
            status = "MATCHED"
            matched = exact
            score = 1.0
            candidates = [{
                "codigo_centro": exact["codigo_cne"],
                "nombre": exact["nombre"],
                "municipio": exact["municipio"],
                "parroquia": exact["parroquia"],
                "confidence_score": 1.0,
            }]
        else:
            fuzzy, score, candidates = _best_fuzzy_match(centro, registry)
            if fuzzy and score >= 0.88:
                status = "MATCHED"
                matched = fuzzy
            elif fuzzy and score >= 0.72:
                status = "AMBIGUOUS"
                matched = None
            else:
                status = "NEW"
                matched = None

        rows.append({
            "row_index": idx,
            "match_status": status,
            "confidence_score": round(score, 3),
            "centro": centro,
            "matched_codigo_centro": matched["codigo_cne"] if matched else None,
            "matched_nombre": matched["nombre"] if matched else None,
            "candidates": candidates,
            "match_hint": hints[idx] if isinstance(hints, list) and idx < len(hints) else None,
        })
        stats[status] += 1
    return {"rows": rows, "stats": stats}


def _match_centros_request(centros: list[dict], hints: Any | None, eleccion_id: int) -> dict[str, Any]:
    db = get_db()
    try:
        ensure_tm_ai_tables(db)
        result = _match_centros(db, centros, hints)
        if eleccion_id:
            total_registry = db.execute("SELECT COUNT(*) c FROM centros").fetchone()["c"]
            linked = db.execute(
                "SELECT COUNT(*) c FROM election_centers WHERE eleccion_id = ?",
                (eleccion_id,),
            ).fetchone()["c"]
            result["registry_not_present_count"] = max(0, total_registry - linked - len(centros))
        return result
    finally:
        db.close()


def _obtener_o_crear_geo(
    db: sqlite3.Connection,
    estado: str | None,
    municipio: str | None,
    parroquia: str | None,
    cod_estado: str | None = None,
    codigo_centro: str | None = None,
) -> tuple[int, int | None, int | None]:
    estado_nom = _canonical_estado_name(estado or "SIN ESTADO") or "SIN ESTADO"
    row = _estado_row_for_centro(db, {"estado": estado, "cod_estado": cod_estado}, codigo_centro)
    if row:
        id_estado = row["id"]
    else:
        code = _estado_code_from_centro({"cod_estado": cod_estado}, codigo_centro)
        if not code:
            code = f"AI{db.execute('SELECT COUNT(*) c FROM estados').fetchone()['c'] + 1:02d}"
        db.execute("INSERT INTO estados (codigo_cne, nombre) VALUES (?, ?)", (code, estado_nom))
        id_estado = db.execute("SELECT last_insert_rowid()").fetchone()[0]

    id_municipio = None
    mun_nom = _normalize_match_text(municipio or "")
    mun_match = _geo_match_name(municipio or "")
    if mun_nom:
        for row in db.execute("SELECT id, nombre FROM municipios WHERE id_estado=?", (id_estado,)).fetchall():
            if _geo_match_name(row["nombre"]) == mun_match:
                id_municipio = row["id"]
                break
        else:
            code = f"AI{db.execute('SELECT COUNT(*) c FROM municipios WHERE id_estado=?', (id_estado,)).fetchone()['c'] + 1:02d}"
            db.execute("INSERT INTO municipios (id_estado, codigo_cne, nombre) VALUES (?, ?, ?)", (id_estado, code, mun_nom))
            id_municipio = db.execute("SELECT last_insert_rowid()").fetchone()[0]

    id_parroquia = None
    parr_nom = _normalize_match_text(parroquia or "")
    parr_match = _geo_match_name(parroquia or "")
    if id_municipio and parr_nom:
        for row in db.execute("SELECT id, nombre FROM parroquias WHERE id_municipio=?", (id_municipio,)).fetchall():
            if _geo_match_name(row["nombre"]) == parr_match:
                id_parroquia = row["id"]
                break
        else:
            code = f"AI{db.execute('SELECT COUNT(*) c FROM parroquias WHERE id_municipio=?', (id_municipio,)).fetchone()['c'] + 1:02d}"
            db.execute("INSERT INTO parroquias (id_municipio, codigo_cne, nombre) VALUES (?, ?, ?)", (id_municipio, code, parr_nom))
            id_parroquia = db.execute("SELECT last_insert_rowid()").fetchone()[0]
    return id_estado, id_municipio, id_parroquia


def _to_int_or_none(value: Any) -> int | None:
    if value in (None, ""):
        return None
    match = re.search(r"\d+", str(value).replace(".", "").replace(",", ""))
    return int(match.group(0)) if match else None


def _confirmed_estado_ids(db: sqlite3.Connection, rows: list[dict]) -> list[int]:
    ids: set[int] = set()
    for item in rows:
        if item.get("match_status") == "EXTRACTION_ERROR":
            continue

        centro = _centro_from_confirm_item(item)
        codigo = (
            item.get("resolved_codigo_centro")
            or item.get("matched_codigo_centro")
            or centro.get("codigo_centro")
            or centro.get("codigo_cne")
        )
        codigo_norm = _normalize_center_code(codigo)

        if codigo_norm:
            existing = db.execute("SELECT id_estado FROM centros WHERE codigo_cne=?", (codigo_norm,)).fetchone()
            if existing and existing["id_estado"] is not None:
                ids.add(existing["id_estado"])

        row = _estado_row_for_centro(db, centro, codigo_norm)
        if row:
            ids.add(row["id"])

        canonical = _canonical_estado_name(centro.get("estado") or "")
        ids.update(_estado_ids_by_canonical(db, canonical))

    return sorted(ids)


def _deactivate_tm_scope(db: sqlite3.Connection, eleccion_id: int, estado_ids: list[int]) -> None:
    if not estado_ids:
        return
    placeholders = ",".join("?" for _ in estado_ids)
    db.execute(f"UPDATE centros SET activo = 0 WHERE id_estado IN ({placeholders})", estado_ids)
    db.execute(
        f"""
        UPDATE election_centers
        SET eligible = 0, updated_at = datetime('now')
        WHERE eleccion_id = ?
          AND centro_id IN (
              SELECT codigo_cne FROM centros WHERE id_estado IN ({placeholders})
          )
        """,
        [eleccion_id, *estado_ids],
    )


def _upsert_ai_center(db: sqlite3.Connection, eleccion_id: int, item: dict[str, Any], source_file: str) -> str:
    centro = item["centro"]
    codigo = item.get("resolved_codigo_centro") or item.get("matched_codigo_centro") or centro.get("codigo_centro")
    codigo = _normalize_center_code(codigo)
    if not codigo:
        codigo = f"AI_{eleccion_id}_{uuid.uuid4().hex[:12]}"

    num_mesas = _to_int_or_none(centro.get("num_mesas"))
    num_electores = _to_int_or_none(centro.get("num_electores"))
    direccion = centro.get("direccion")
    existing = db.execute("SELECT * FROM centros WHERE codigo_cne=?", (codigo,)).fetchone()

    if existing:
        db.execute("""
            UPDATE centros SET
                num_mesas = COALESCE(?, num_mesas),
                num_electores = COALESCE(?, num_electores),
                direccion = CASE
                    WHEN (direccion IS NULL OR TRIM(direccion) = '') AND ? IS NOT NULL AND TRIM(?) != ''
                    THEN ?
                    ELSE direccion
                END,
                activo = 1
            WHERE codigo_cne = ?
        """, (num_mesas, num_electores, direccion, direccion or "", direccion, codigo))
    else:
        id_estado, id_municipio, id_parroquia = _obtener_o_crear_geo(
            db,
            centro.get("estado"),
            centro.get("municipio"),
            centro.get("parroquia"),
            centro.get("cod_estado"),
            codigo,
        )
        db.execute("""
            INSERT INTO centros (
                codigo_cne, nombre, direccion, id_parroquia, id_municipio, id_estado,
                num_mesas, num_electores, activo
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, 1)
        """, (
            codigo,
            _normalize_match_text(centro.get("nombre_centro") or codigo),
            direccion,
            id_parroquia, id_municipio, id_estado,
            num_mesas or 0, num_electores or 0,
        ))

    db.execute("""
        INSERT INTO election_centers (eleccion_id, centro_id, eligible, source_file, campos_extra, updated_at)
        VALUES (?, ?, 1, ?, ?, datetime('now'))
        ON CONFLICT(eleccion_id, centro_id) DO UPDATE SET
            eligible=excluded.eligible,
            source_file=excluded.source_file,
            campos_extra=excluded.campos_extra,
            updated_at=datetime('now')
    """, (
        eleccion_id, codigo, source_file,
        json.dumps(centro.get("campos_extra") or {}, ensure_ascii=False),
    ))
    return codigo


@app.post("/api/tm/ai-extract")
async def tm_ai_extract(request: Request, archivo: UploadFile | None = File(None)):
    content_type = request.headers.get("content-type", "")
    if "multipart/form-data" in content_type:
        form = await request.form()
        eleccion_id = int(form.get("eleccion_id") or 0)
        source_file = str(form.get("source_file") or (archivo.filename if archivo else "archivo"))
        text = str(form.get("text") or "")
        if archivo and not text:
            ext = Path(archivo.filename or "").suffix.lower()
            if ext not in (".pdf", ".xlsx", ".xls", ".xlsm", ".csv", ".docx", ".txt"):
                raise HTTPException(400, "Formato no soportado")
            data = await archivo.read()
            if len(data) > 50 * 1024 * 1024:
                raise HTTPException(400, "Archivo mayor a 50MB")
            tmp_path = UPLOAD_DIR / f"tm_ai_{uuid.uuid4().hex}{ext}"
            tmp_path.write_bytes(data)
            try:
                text = await asyncio.to_thread(_read_upload_text, tmp_path, ext)
            finally:
                tmp_path.unlink(missing_ok=True)
    else:
        payload = await request.json()
        eleccion_id = int(payload.get("eleccion_id") or 0)
        source_file = str(payload.get("source_file") or "texto")
        text = str(payload.get("text") or "")

    if not eleccion_id:
        raise HTTPException(400, "eleccion_id requerido")
    if not text.strip():
        raise HTTPException(400, "No se pudo extraer texto del archivo")

    chunks = _chunk_text(text, TM_AI_CHUNK_SIZE)
    merged = {"detected_columns": [], "field_notes": {}, "centros": [], "match_hints": []}
    sys.path.insert(0, str(BASE_DIR))
    import agent

    db = get_db()
    try:
        cfg_row = get_ai_config(db)
        provider = cfg_row["provider"]
        cfg = dict(cfg_row)
        cfg["max_tokens"] = max(int(cfg.get("max_tokens") or 300), 12000)
    finally:
        db.close()

    for idx, chunk in enumerate(chunks, 1):
        user_prompt = (
            f"Source file: {source_file}\n"
            f"Chunk {idx} of {len(chunks)}. Preserve source column names.\n\n"
            f"RAW EXTRACTED TEXT:\n{chunk}"
        )
        last_error = None
        for attempt in range(TM_AI_MAX_ATTEMPTS):
            try:
                raw = await agent.ask_structured_async(
                    TM_AI_SYSTEM_PROMPT,
                    user_prompt,
                    provider,
                    cfg,
                )
                parsed = _extract_json_object(raw)
                merged["detected_columns"].extend(parsed.get("detected_columns") or [])
                merged["field_notes"].update(parsed.get("field_notes") or {})
                merged["centros"].extend(parsed.get("centros") or [])
                merged["match_hints"].extend(parsed.get("match_hints") or [])
                break
            except Exception as exc:
                last_error = exc
                if attempt < TM_AI_MAX_ATTEMPTS - 1:
                    await asyncio.sleep(_ai_retry_delay(attempt, exc))
        else:
            merged["field_notes"][f"chunk_{idx}_error"] = str(last_error)

    merged["detected_columns"] = list(dict.fromkeys(str(c) for c in merged["detected_columns"]))
    if not merged["centros"]:
        return JSONResponse({
            "ok": False,
            "source_file": source_file,
            "raw_text_sample": text[:2000],
            **merged,
        }, status_code=422)
    return JSONResponse({"ok": True, "source_file": source_file, **merged})


@app.post("/api/tm/fuzzy-match")
async def tm_fuzzy_match(request: Request):
    payload = await request.json()
    centros = payload.get("centros") or []
    hints = payload.get("match_hints") or []
    eleccion_id = int(payload.get("eleccion_id") or 0)
    result = await asyncio.to_thread(_match_centros_request, centros, hints, eleccion_id)
    return JSONResponse(result)


@app.post("/api/tm/confirm")
async def tm_confirm(request: Request):
    payload = await request.json()
    eleccion_id = int(payload.get("eleccion_id") or 0)
    rows = payload.get("rows") or []
    source_files = payload.get("source_files") or []
    dry_run = bool(payload.get("dry_run"))
    if not eleccion_id:
        raise HTTPException(400, "eleccion_id requerido")

    blockers = [r for r in rows if r.get("match_status") in ("AMBIGUOUS", "CONFLICT") and not r.get("resolved_codigo_centro")]
    if blockers:
        raise HTTPException(400, "Hay filas AMBIGUOUS/CONFLICT sin resolver")

    db = get_db()
    stats = {"MATCHED": 0, "NEW": 0, "AMBIGUOUS": 0, "CONFLICT": 0, "EXTRACTION_ERROR": 0, "written": 0}
    detected_columns = payload.get("detected_columns") or []
    field_notes = payload.get("field_notes") or {}
    try:
        ensure_tm_ai_tables(db)
        affected_estado_ids = _confirmed_estado_ids(db, rows)
        if not dry_run:
            db.execute("BEGIN IMMEDIATE")
            _deactivate_tm_scope(db, eleccion_id, affected_estado_ids)

        for item in rows:
            status = item.get("match_status") or "EXTRACTION_ERROR"
            stats[status] = stats.get(status, 0) + 1
            if status == "EXTRACTION_ERROR":
                continue
            if not dry_run:
                _upsert_ai_center(db, eleccion_id, item, item.get("source_file") or ", ".join(source_files))
            stats["written"] += 1

        if not dry_run:
            db.execute("""
                INSERT INTO tm_ingestion_logs
                    (eleccion_id, source_files, detected_columns, field_notes, match_stats, user)
                VALUES (?, ?, ?, ?, ?, ?)
            """, (
                eleccion_id,
                json.dumps(source_files, ensure_ascii=False),
                json.dumps(detected_columns, ensure_ascii=False),
                json.dumps(field_notes, ensure_ascii=False),
                json.dumps(stats, ensure_ascii=False),
                payload.get("user") or "local",
            ))
            db.commit()
        return JSONResponse({
            "ok": True,
            "dry_run": dry_run,
            "stats": stats,
            "replacement_scope": {"estado_ids": affected_estado_ids},
        })
    except Exception:
        db.rollback()
        raise
    finally:
        db.close()


# ══════════════════════════════════════════════════════════════════
# MUESTRA — Selección de centros para el exit poll
# ══════════════════════════════════════════════════════════════════

@app.get("/muestra", response_class=HTMLResponse)
async def muestra_index(
    request: Request,
    msg: str = "",
    cat: str = "success",
    q: str = "",
    estado: str = "",
    municipio: str = "",
    parroquia: str = "",
    estatus: str = "",
    clasificacion: str = "",
    offset: int = 0,
):
    db = get_db()
    eleccion = db.execute("SELECT * FROM elecciones WHERE activa=1 LIMIT 1").fetchone()
    sys.path.insert(0, str(BASE_DIR))
    import muestra_lab

    laboratorio = muestra_lab.construir_laboratorio(
        db,
        dict(eleccion) if eleccion else None,
        q=q,
        estado=estado,
        municipio=municipio,
        parroquia=parroquia,
        estatus=estatus,
        clasificacion=clasificacion,
        limit=300,
        offset=offset,
    )

    eleccion_ref = _eleccion_ref_referencia(db)
    muestra_actual = []
    if eleccion:
        muestra_actual = db.execute(
            """SELECT m.id, m.codigo_centro, ct.nombre as centro_nombre,
                      e.nombre as estado, mu.nombre as municipio, p.nombre as parroquia,
                      m.tipo_centro, m.motivo, m.score_snapshot, m.confianza_snapshot,
                      ct.num_electores, ct.num_mesas,
                      rh.pct_oposicion, rh.pct_gobierno
               FROM muestra m
               JOIN centros ct ON m.codigo_centro=ct.codigo_cne
               JOIN estados e ON ct.id_estado=e.id
               LEFT JOIN municipios mu ON ct.id_municipio=mu.id
               LEFT JOIN parroquias p ON ct.id_parroquia=p.id
               LEFT JOIN resultados_historicos rh
                   ON rh.codigo_centro=m.codigo_centro AND rh.eleccion_ref=?
               WHERE m.id_eleccion=? AND m.activo=1
               ORDER BY e.nombre, mu.nombre, ct.num_electores DESC""",
            (eleccion_ref, eleccion["id"])
        ).fetchall()

    # Resultado nacional de referencia
    nac = db.execute("""
        SELECT SUM(votos_validos) v, SUM(votos_gobierno) g, SUM(votos_oposicion) o
        FROM resultados_historicos WHERE eleccion_ref=?
    """, (eleccion_ref,)).fetchone()
    pct_nac = {}
    if nac and nac["v"]:
        pct_nac["gobierno"] = round(100 * nac["g"] / nac["v"], 1)
        pct_nac["oposicion"] = round(100 * nac["o"] / nac["v"], 1)

    # Refs historicas disponibles
    refs = db.execute(
        "SELECT DISTINCT eleccion_ref FROM resultados_historicos"
    ).fetchall()

    db.close()
    return templates.TemplateResponse(request=request, name="muestra.html", context={
        "eleccion": eleccion,
        "muestra": muestra_actual, "pct_nac": pct_nac,
        "refs": refs, "eleccion_ref": eleccion_ref,
        "msg": msg, "cat": cat,
        "lab": laboratorio,
        "filtros": {
            "q": q,
            "estado": estado,
            "municipio": municipio,
            "parroquia": parroquia,
            "estatus": estatus,
            "clasificacion": clasificacion,
            "offset": offset,
        },
    })


@app.post("/muestra/agregar")
async def muestra_agregar(request: Request):
    form = await request.form()
    codigo = str(form.get("codigo_centro") or "").strip()
    motivo = str(form.get("motivo") or "Seleccion manual desde laboratorio").strip()
    db = get_db()
    eleccion = db.execute("SELECT * FROM elecciones WHERE activa=1 LIMIT 1").fetchone()
    if not eleccion:
        db.close()
        return RedirectResponse("/muestra?msg=Active+una+eleccion+primero&cat=warning", status_code=303)
    sys.path.insert(0, str(BASE_DIR))
    import muestra_lab

    ok = muestra_lab.agregar_centro(db, eleccion["id"], codigo, motivo=motivo, usuario="local")
    db.close()
    if not ok:
        return RedirectResponse("/muestra?msg=Centro+no+seleccionable+o+sin+registro+activo&cat=warning", status_code=303)
    return RedirectResponse("/muestra?msg=Centro+agregado+a+la+muestra&cat=success", status_code=303)


@app.get("/muestra/generar", response_class=HTMLResponse)
async def muestra_generar(
    request: Request,
    centros_por_unidad: int = 2,
    candidatos_por_unidad: int = 5,
    umbral_pct: float = 10.0,
    eleccion_ref: str = "",
):
    db = get_db()
    eleccion = db.execute("SELECT * FROM elecciones WHERE activa=1 LIMIT 1").fetchone()
    if not eleccion:
        db.close()
        return RedirectResponse("/muestra?msg=Active+una+elección+primero&cat=warning", status_code=303)
    eleccion_ref = eleccion_ref or _eleccion_ref_referencia(db) or "2024-presidencial"

    sys.path.insert(0, str(BASE_DIR))
    import selector_muestra
    # Reload in case module was cached
    import importlib
    importlib.reload(selector_muestra)

    nac = selector_muestra.resultado_nacional(db, eleccion_ref)
    candidatos = selector_muestra.generar_candidatos(
        db, id_eleccion=eleccion["id"], eleccion_ref=eleccion_ref,
        candidatos_por_unidad=candidatos_por_unidad,
        umbral_pct=umbral_pct,
    )

    db.close()
    return templates.TemplateResponse(request=request, name="muestra_generar.html", context={
        "eleccion": eleccion,
        "candidatos": candidatos, "nac": nac,
        "centros_por_unidad": centros_por_unidad,
        "candidatos_por_unidad": candidatos_por_unidad,
        "umbral_pct": umbral_pct,
        "eleccion_ref": eleccion_ref,
    })


@app.post("/muestra/aplicar")
async def muestra_aplicar(request: Request):
    form = await request.form()
    codigos = form.getlist("codigo_centro")
    id_eleccion = int(form.get("id_eleccion", 0))

    if not codigos or not id_eleccion:
        return RedirectResponse("/muestra?msg=No+se+seleccionaron+centros&cat=warning", status_code=303)

    db = get_db()
    sys.path.insert(0, str(BASE_DIR))
    import selector_muestra
    import importlib
    importlib.reload(selector_muestra)

    n = selector_muestra.aplicar_muestra(db, id_eleccion, codigos)
    db.close()
    return RedirectResponse(f"/muestra?msg=Muestra+creada+con+{n}+centros&cat=success", status_code=303)


@app.post("/muestra/{mid}/quitar")
async def muestra_quitar(mid: int):
    db = get_db()
    db.execute("DELETE FROM pesos WHERE id_muestra=?", (mid,))
    db.execute("DELETE FROM muestra WHERE id=?", (mid,))
    db.commit()
    db.close()
    return RedirectResponse("/muestra?msg=Centro+removido+de+la+muestra&cat=warning", status_code=303)


# ══════════════════════════════════════════════════════════════════
# VISUALIZACIÓN — Heatmap y Dashboard interactivo
# ══════════════════════════════════════════════════════════════════

VIZ_DIR = BASE_DIR / "static" / "viz"
VIZ_DIR.mkdir(parents=True, exist_ok=True)


def _eleccion_ref_referencia(db) -> str | None:
    """Ref histórica más reciente (los refs empiezan por año, ordenan lexicográficamente)."""
    row = db.execute(
        "SELECT eleccion_ref FROM resultados_historicos ORDER BY eleccion_ref DESC LIMIT 1"
    ).fetchone()
    return row["eleccion_ref"] if row else None


def _datos_ventaja_por_estado(db, eleccion_ref: str) -> dict:
    """Extrae ventaja gobierno-oposición por estado desde resultados_historicos."""
    rows = db.execute("""
        SELECT e.nombre,
               SUM(rh.votos_gobierno) as gob,
               SUM(rh.votos_oposicion) as opo,
               SUM(rh.votos_validos) as val
        FROM resultados_historicos rh
        JOIN centros c ON rh.codigo_centro = c.codigo_cne
        JOIN estados e ON c.id_estado = e.id
        WHERE c.activo = 1 AND rh.eleccion_ref = ?
        GROUP BY e.id
    """, (eleccion_ref,)).fetchall()
    datos = {}
    for r in rows:
        if r["val"] and r["val"] > 0:
            pct_gob = 100 * r["gob"] / r["val"]
            pct_opo = 100 * r["opo"] / r["val"]
            ventaja = round(pct_gob - pct_opo, 1)
            # Normalizar nombre: "EDO. ZULIA" → "Zulia", "DTTO. CAPITAL" → "Distrito Capital"
            nombre = r["nombre"]
            nombre = nombre.replace("EDO. ", "").replace("DTTO. ", "Distrito ")
            nombre = nombre.replace("DELTA AMAC", "Delta Amacuro")
            nombre = nombre.replace("LA GUAIRA", "La Guaira")
            nombre = nombre.replace("NVA. ESPARTA", "Nueva Esparta")
            if nombre not in ("Distrito Capital", "La Guaira", "Delta Amacuro", "Nueva Esparta"):
                nombre = nombre.title()
            datos[_norm_estado(r["nombre"])] = ventaja
    return datos


def _datos_ventaja_muestra(db, id_eleccion: int, eleccion_ref: str) -> dict:
    """Ventaja solo para centros en la muestra, por estado."""
    rows = db.execute("""
        SELECT e.nombre,
               SUM(rh.votos_gobierno) as gob,
               SUM(rh.votos_oposicion) as opo,
               SUM(rh.votos_validos) as val
        FROM muestra m
        JOIN centros c ON m.codigo_centro = c.codigo_cne
        JOIN estados e ON c.id_estado = e.id
        LEFT JOIN resultados_historicos rh
            ON rh.codigo_centro = m.codigo_centro AND rh.eleccion_ref = ?
        WHERE m.id_eleccion = ? AND m.activo = 1
        GROUP BY e.id
    """, (eleccion_ref, id_eleccion)).fetchall()
    datos = {}
    for r in rows:
        if r["val"] and r["val"] > 0:
            pct_gob = 100 * r["gob"] / r["val"]
            pct_opo = 100 * r["opo"] / r["val"]
            ventaja = round(pct_gob - pct_opo, 1)
            nombre = r["nombre"]
            nombre = nombre.replace("EDO. ", "").replace("DTTO. ", "Distrito ")
            nombre = nombre.replace("DELTA AMAC", "Delta Amacuro")
            nombre = nombre.replace("LA GUAIRA", "La Guaira")
            nombre = nombre.replace("NVA. ESPARTA", "Nueva Esparta")
            if nombre not in ("Distrito Capital", "La Guaira", "Delta Amacuro", "Nueva Esparta"):
                nombre = nombre.title()
            datos[_norm_estado(r["nombre"])] = ventaja
    return datos


def _norm_municipio(nombre: str) -> str:
    """Normaliza nombres CNE de municipio para cruzarlos con ADM2."""
    nombre = (nombre or "").strip()
    for pref in ("MP. ", "CM. ", "MCPIO. ", "MUNICIPIO "):
        if nombre.upper().startswith(pref):
            nombre = nombre[len(pref):]
            break
    especiales = {
        "BLVNO LIBERTADOR": "Libertador",
        "LIBERTADOR": "Libertador",
    }
    return especiales.get(nombre.upper(), nombre.title())


def _datos_ventaja_por_municipio(db, eleccion_ref: str) -> dict:
    """Extrae ventaja gobierno-oposicion por municipio."""
    rows = db.execute("""
        SELECT e.nombre AS estado, mu.nombre AS municipio,
               SUM(rh.votos_gobierno) AS gob,
               SUM(rh.votos_oposicion) AS opo,
               SUM(rh.votos_validos) AS val
        FROM resultados_historicos rh
        JOIN centros c ON rh.codigo_centro = c.codigo_cne
        JOIN estados e ON c.id_estado = e.id
        JOIN municipios mu ON c.id_municipio = mu.id
        WHERE c.activo = 1 AND rh.eleccion_ref = ?
        GROUP BY e.id, mu.id
    """, (eleccion_ref,)).fetchall()
    datos = {}
    for r in rows:
        if r["val"] and r["val"] > 0:
            pct_gob = 100 * r["gob"] / r["val"]
            pct_opo = 100 * r["opo"] / r["val"]
            datos[(_norm_estado(r["estado"]), _norm_municipio(r["municipio"]))] = round(pct_gob - pct_opo, 1)
    return datos


def _datos_ventaja_muestra_municipio(db, id_eleccion: int, eleccion_ref: str) -> dict:
    """Extrae ventaja de centros en muestra por municipio."""
    rows = db.execute("""
        SELECT e.nombre AS estado, mu.nombre AS municipio,
               SUM(rh.votos_gobierno) AS gob,
               SUM(rh.votos_oposicion) AS opo,
               SUM(rh.votos_validos) AS val
        FROM muestra m
        JOIN centros c ON m.codigo_centro = c.codigo_cne
        JOIN estados e ON c.id_estado = e.id
        JOIN municipios mu ON c.id_municipio = mu.id
        LEFT JOIN resultados_historicos rh
            ON rh.codigo_centro = m.codigo_centro AND rh.eleccion_ref = ?
        WHERE m.id_eleccion = ? AND m.activo = 1
        GROUP BY e.id, mu.id
    """, (eleccion_ref, id_eleccion)).fetchall()
    datos = {}
    for r in rows:
        if r["val"] and r["val"] > 0:
            pct_gob = 100 * r["gob"] / r["val"]
            pct_opo = 100 * r["opo"] / r["val"]
            datos[(_norm_estado(r["estado"]), _norm_municipio(r["municipio"]))] = round(pct_gob - pct_opo, 1)
    return datos


def _tendencia_simulada(datos_ventaja: dict, n_puntos: int = 15) -> dict:
    """Genera tendencia simulada a partir de datos de ventaja históricos.
    Simula la llegada progresiva de datos desde las 7AM con la tendencia
    convergiendo al resultado final conocido."""
    import datetime
    import random
    tendencias = {}
    start = datetime.datetime(2025, 1, 1, 7, 0)

    # Nacional: promedio ponderado
    if datos_ventaja:
        avg_ventaja = sum(datos_ventaja.values()) / len(datos_ventaja)
        final_gob = 50 + avg_ventaja / 2
        final_opo = 50 - avg_ventaja / 2
    else:
        final_gob, final_opo = 50, 50

    etiquetas = {
        (f"{k[0]} - {k[1]}" if isinstance(k, tuple) else str(k)): v
        for k, v in datos_ventaja.items()
    }

    for nombre in ["VENEZUELA"] + [n.upper() for n in etiquetas.keys()]:
        if nombre == "VENEZUELA":
            tgt_gob, tgt_opo = final_gob, final_opo
        else:
            orig = next((k for k in etiquetas if k.upper() == nombre), None)
            v = etiquetas.get(orig, 0) if orig else 0
            tgt_gob = 50 + v / 2
            tgt_opo = 50 - v / 2

        puntos = []
        gob = 50 + random.uniform(-5, 5)  # Inicia cerca de 50/50
        for i in range(n_puntos):
            hora = (start + datetime.timedelta(minutes=40 * i)).strftime('%H:%M')
            # Converge al resultado final con ruido decreciente
            t = (i + 1) / n_puntos
            noise = random.uniform(-3, 3) * (1 - t)
            gob = tgt_gob * t + gob * (1 - t) + noise
            gob = max(5, min(95, gob))
            opo = 100 - gob
            puntos.append({'hora': hora, 'gob': round(gob, 1), 'opo': round(opo, 1)})
        tendencias[nombre] = puntos

    return tendencias


def _total_referencia_dashboard(
    db, eleccion_ref: str, id_eleccion: int | None = None, fuente: str = "todos"
) -> int:
    """Cuenta opiniones de referencia con la misma fuente usada por /visualizacion."""
    if fuente == "muestra" and id_eleccion:
        row = db.execute("""
            SELECT COALESCE(SUM(rh.votos_validos), 0) AS total
            FROM muestra m
            LEFT JOIN resultados_historicos rh
                ON rh.codigo_centro = m.codigo_centro AND rh.eleccion_ref = ?
            WHERE m.id_eleccion = ? AND m.activo = 1
        """, (eleccion_ref, id_eleccion)).fetchone()
    else:
        row = db.execute("""
            SELECT COALESCE(SUM(rh.votos_validos), 0) AS total
            FROM resultados_historicos rh
            JOIN centros c ON rh.codigo_centro = c.codigo_cne
            WHERE c.activo = 1 AND rh.eleccion_ref = ?
        """, (eleccion_ref,)).fetchone()
    return int(row["total"] or 0)


def _datos_dashboard_referencia(
    db,
    id_eleccion: int | None,
    nivel: str = "estado",
    fuente: str = "todos",
) -> tuple[dict, dict, int]:
    """Devuelve la misma data base que usa el dashboard estatico de /visualizacion."""
    eleccion_ref = _eleccion_ref_referencia(db)
    if not eleccion_ref:
        return {}, {}, 0

    if nivel == "municipio":
        if fuente == "muestra" and id_eleccion:
            datos_ventaja = _datos_ventaja_muestra_municipio(db, id_eleccion, eleccion_ref)
        else:
            datos_ventaja = _datos_ventaja_por_municipio(db, eleccion_ref)
    elif fuente == "muestra" and id_eleccion:
        datos_ventaja = _datos_ventaja_muestra(db, id_eleccion, eleccion_ref)
    else:
        datos_ventaja = _datos_ventaja_por_estado(db, eleccion_ref)

    if not datos_ventaja:
        return {}, {}, 0
    return (
        datos_ventaja,
        _tendencia_simulada(datos_ventaja),
        _total_referencia_dashboard(db, eleccion_ref, id_eleccion, fuente),
    )


def _nombres_candidatos(db, id_eleccion: int = None) -> dict:
    """Obtiene nombres de candidatos gobierno/oposición de la elección."""
    cand = {'gobierno': 'Gobierno', 'oposicion': 'Oposici\u00f3n'}
    if id_eleccion:
        for bando in ('gobierno', 'oposicion'):
            row = db.execute(
                "SELECT nombre FROM candidatos WHERE id_eleccion=? AND bando=? ORDER BY orden LIMIT 1",
                (id_eleccion, bando)
            ).fetchone()
            if row:
                cand[bando] = row["nombre"]
    return cand


@app.get("/visualizacion", response_class=HTMLResponse)
async def visualizacion_index(request: Request, msg: str = "", cat: str = "success"):
    db = get_db()
    eleccion = db.execute("SELECT * FROM elecciones WHERE activa=1 LIMIT 1").fetchone()

    tiene_resultados = db.execute(
        "SELECT COUNT(*) c FROM resultados_historicos"
    ).fetchone()["c"] > 0

    tiene_muestra = False
    if eleccion:
        tiene_muestra = db.execute(
            "SELECT COUNT(*) c FROM muestra WHERE id_eleccion=?", (eleccion["id"],)
        ).fetchone()["c"] > 0

    # Verificar si ya hay archivos generados
    heatmap_existe = (VIZ_DIR / "heatmap.html").exists()
    dashboard_existe = (VIZ_DIR / "dashboard.html").exists()

    refs = db.execute(
        "SELECT DISTINCT eleccion_ref FROM resultados_historicos"
    ).fetchall()

    db.close()
    return templates.TemplateResponse(request=request, name="visualizacion.html", context={
        "eleccion": eleccion,
        "tiene_resultados": tiene_resultados,
        "tiene_muestra": tiene_muestra,
        "heatmap_existe": heatmap_existe,
        "dashboard_existe": dashboard_existe,
        "refs": refs,
        "msg": msg, "cat": cat,
    })


@app.post("/visualizacion/generar")
async def visualizacion_generar(
    request: Request,
    tipo: str = Form("heatmap"),
    nivel: str = Form("estado"),
    fuente: str = Form("todos"),
):
    db = get_db()
    eleccion = db.execute("SELECT * FROM elecciones WHERE activa=1 LIMIT 1").fetchone()
    eid = eleccion["id"] if eleccion else None

    candidatos_dict = _nombres_candidatos(db, eid)

    # Obtener datos de ventaja (siempre de una sola eleccion_ref)
    eleccion_ref = _eleccion_ref_referencia(db)
    if not eleccion_ref:
        datos_ventaja = {}
    elif nivel == "municipio":
        if fuente == "muestra" and eid:
            datos_ventaja = _datos_ventaja_muestra_municipio(db, eid, eleccion_ref)
        else:
            datos_ventaja = _datos_ventaja_por_municipio(db, eleccion_ref)
    elif fuente == "muestra" and eid:
        datos_ventaja = _datos_ventaja_muestra(db, eid, eleccion_ref)
    else:
        datos_ventaja = _datos_ventaja_por_estado(db, eleccion_ref)

    db.close()

    if not datos_ventaja:
        return RedirectResponse(
            "/visualizacion?msg=No+hay+datos+de+resultados+hist%C3%B3ricos&cat=warning",
            status_code=303
        )

    sys.path.insert(0, str(BASE_DIR))

    titulo = eleccion["nombre"] if eleccion else "Exit Poll Venezuela"

    try:
        if tipo == "dashboard":
            import generador_dashboard
            import importlib
            importlib.reload(generador_dashboard)

            tendencias = _tendencia_simulada(datos_ventaja)
            ruta = str(VIZ_DIR / "dashboard.html")
            generador_dashboard.generar_dashboard(
                datos_ventaja, tendencias,
                nivel=nivel, ruta_salida=ruta,
                titulo=titulo, candidatos=candidatos_dict
            )
            return RedirectResponse(
                "/visualizacion?msg=Dashboard+generado+exitosamente&cat=success",
                status_code=303
            )
        else:
            import generador_heatmap
            import importlib
            importlib.reload(generador_heatmap)

            ruta = str(VIZ_DIR / "heatmap.html")
            generador_heatmap.generar_heatmap(
                datos_ventaja, nivel=nivel, ruta_salida=ruta,
                titulo=titulo, candidatos=candidatos_dict
            )
            return RedirectResponse(
                "/visualizacion?msg=Heatmap+generado+exitosamente&cat=success",
                status_code=303
            )
    except Exception as e:
        import traceback
        traceback.print_exc()
        error_msg = str(e)[:200]
        return RedirectResponse(
            f"/visualizacion?msg=Error:+{error_msg}&cat=danger",
            status_code=303
        )


# ══════════════════════════════════════════════════════════════════
# LIVE — Dashboard en tiempo real para el Showcase
# ══════════════════════════════════════════════════════════════════

def _norm_estado(nombre: str) -> str:
    """Normaliza nombres de estado del CNE para que coincidan con el GeoJSON."""
    nombre = (nombre
              .replace("EDO. ", "")
              .replace("DTTO. ", "Distrito ")
              .replace("DELTA AMAC", "Delta Amacuro")
              .replace("LA GUAIRA", "La Guaira")
              .replace("NVA. ESPARTA", "Nueva Esparta"))
    if nombre not in ("Distrito Capital", "La Guaira", "Delta Amacuro", "Nueva Esparta"):
        nombre = nombre.title()
    return nombre


def _datos_vivos(db, id_eleccion: int) -> tuple:
    """
    Lee opiniones reales de la BD y devuelve
    (datos_ventaja, datos_tendencia, total_votos).

    datos_ventaja  : {nombre_estado: ventaja_float}   — gobierno - oposición en %
    datos_tendencia: {NOMBRE_UPPER: [{"hora","gob","opo"}, ...]}  — acumulados por turno
    """
    total = db.execute("""
        SELECT COUNT(*) c FROM votos v
        JOIN muestra m ON m.codigo_centro = v.codigo_centro
        WHERE m.id_eleccion = ? AND v.valido = 1
    """, (id_eleccion,)).fetchone()["c"]

    if total == 0:
        return {}, {}, 0

    # ── Ventaja final por estado ──────────────────────────────────
    rows_est = db.execute("""
        SELECT
            est.nombre                                              AS estado,
            SUM(CASE WHEN ca.bando = 'gobierno'  THEN 1 ELSE 0 END) AS gob,
            SUM(CASE WHEN ca.bando = 'oposicion' THEN 1 ELSE 0 END) AS opo,
            COUNT(*)                                                AS total
        FROM votos v
        JOIN muestra    m   ON m.codigo_centro = v.codigo_centro
        JOIN centros    c   ON c.codigo_cne    = v.codigo_centro
        JOIN estados    est ON est.id          = c.id_estado
        JOIN candidatos ca  ON ca.id           = v.id_candidato
        WHERE v.valido = 1 AND m.id_eleccion = ?
        GROUP BY est.id
    """, (id_eleccion,)).fetchall()

    datos_ventaja = {}
    for r in rows_est:
        if r["total"] > 0:
            ventaja = round(100 * r["gob"] / r["total"] - 100 * r["opo"] / r["total"], 1)
            datos_ventaja[_norm_estado(r["estado"])] = ventaja

    # ── Tendencias acumuladas — nacional ─────────────────────────
    rows_nac = db.execute("""
        SELECT
            v.turno,
            MIN(v.hora)                                             AS hora_min,
            SUM(CASE WHEN ca.bando = 'gobierno'  THEN 1 ELSE 0 END) AS gob,
            SUM(CASE WHEN ca.bando = 'oposicion' THEN 1 ELSE 0 END) AS opo
        FROM votos v
        JOIN muestra    m  ON m.codigo_centro = v.codigo_centro
        JOIN candidatos ca ON ca.id           = v.id_candidato
        WHERE v.valido = 1 AND m.id_eleccion = ?
        GROUP BY v.turno
        ORDER BY v.turno
    """, (id_eleccion,)).fetchall()

    datos_tendencia = {}
    puntos_nac, cum_g, cum_o = [], 0, 0
    for r in rows_nac:
        cum_g += r["gob"]
        cum_o += r["opo"]
        tot = cum_g + cum_o
        if tot:
            h = r["hora_min"]
            hora = h[11:16] if h and "T" in h else (h or "07:00")[:5]
            puntos_nac.append({
                "hora": hora,
                "gob": round(100 * cum_g / tot, 1),
                "opo": round(100 * cum_o / tot, 1),
            })
    datos_tendencia["VENEZUELA"] = puntos_nac

    # ── Tendencias acumuladas — por estado ────────────────────────
    rows_t = db.execute("""
        SELECT
            est.nombre                                              AS estado,
            v.turno,
            MIN(v.hora)                                             AS hora_min,
            SUM(CASE WHEN ca.bando = 'gobierno'  THEN 1 ELSE 0 END) AS gob,
            SUM(CASE WHEN ca.bando = 'oposicion' THEN 1 ELSE 0 END) AS opo
        FROM votos v
        JOIN muestra    m   ON m.codigo_centro = v.codigo_centro
        JOIN centros    c   ON c.codigo_cne    = v.codigo_centro
        JOIN estados    est ON est.id          = c.id_estado
        JOIN candidatos ca  ON ca.id           = v.id_candidato
        WHERE v.valido = 1 AND m.id_eleccion = ?
        GROUP BY est.id, v.turno
        ORDER BY est.nombre, v.turno
    """, (id_eleccion,)).fetchall()

    por_estado = {}
    for r in rows_t:
        por_estado.setdefault(r["estado"], []).append(r)

    for estado_db, turnos in por_estado.items():
        nombre_up = _norm_estado(estado_db).upper()
        puntos, cum_g, cum_o = [], 0, 0
        for r in sorted(turnos, key=lambda x: x["turno"]):
            cum_g += r["gob"]
            cum_o += r["opo"]
            tot = cum_g + cum_o
            if tot:
                h = r["hora_min"]
                hora = h[11:16] if h and "T" in h else (h or "07:00")[:5]
                puntos.append({
                    "hora": hora,
                    "gob": round(100 * cum_g / tot, 1),
                    "opo": round(100 * cum_o / tot, 1),
                })
        datos_tendencia[nombre_up] = puntos

    return datos_ventaja, datos_tendencia, total


def _dashboard_stream_payload(db: sqlite3.Connection) -> dict:
    eleccion = db.execute(
        "SELECT * FROM elecciones WHERE activa = 1 LIMIT 1"
    ).fetchone()
    if not eleccion:
        return {
            "ok": False,
            "reason": "No hay eleccion activa",
            "geo": {},
            "series": {},
            "total_votos": 0,
            "total_opiniones": 0,
        }

    datos_ventaja, datos_tendencia, total_votos = _datos_vivos(db, eleccion["id"])
    fuente_datos = "live"
    if total_votos == 0:
        datos_ventaja, datos_tendencia, total_votos = _datos_dashboard_referencia(
            db, eleccion["id"], nivel="estado", fuente="todos"
        )
        fuente_datos = "dashboard_referencia" if total_votos else "sin_datos"

    return {
        "ok": True,
        "eleccion": eleccion["nombre"],
        "fuente_datos": fuente_datos,
        "geo": datos_ventaja,
        "series": datos_tendencia,
        "total_votos": total_votos,
        "total_opiniones": total_votos,
    }


@app.get("/stream/dashboard")
async def stream_dashboard():
    async def events():
        while True:
            db = get_db()
            try:
                payload = _dashboard_stream_payload(db)
            finally:
                db.close()
            yield f"data: {json.dumps(payload, ensure_ascii=False)}\n\n"
            await asyncio.sleep(60)

    return StreamingResponse(events(), media_type="text/event-stream")


def _html_sin_datos(motivo: str, refresh: int) -> str:
    """Página con mapa base en gris y mensaje de espera."""
    try:
        sys.path.insert(0, str(BASE_DIR))
        import generador_heatmap
        import importlib
        importlib.reload(generador_heatmap)
        tmp = tempfile.mktemp(suffix=".html")
        generador_heatmap.generar_heatmap(
            {}, nivel="estado", ruta_salida=tmp,
            titulo="Exit Poll — En Vivo"
        )
        with open(tmp, encoding="utf-8") as f:
            base_html = f.read()
        os.unlink(tmp)
    except Exception:
        base_html = "<html><head></head><body></body></html>"

    overlay = f"""
<div style="position:fixed;top:50%;left:50%;transform:translate(-50%,-50%);
     z-index:9999;background:rgba(0,0,0,.75);color:white;
     font-family:sans-serif;font-size:16px;padding:28px 40px;
     border-radius:12px;text-align:center;max-width:440px;
     box-shadow:0 4px 24px rgba(0,0,0,.4);">
  <div style="font-size:36px;margin-bottom:14px;">&#x1F4E1;</div>
  <div style="font-weight:bold;margin-bottom:10px;line-height:1.4">{motivo}</div>
  <div style="font-size:12px;opacity:.65;margin-top:6px;">
    Esperando datos en vivo por SSE
  </div>
</div>"""

    html = base_html
    html = re.sub(r"(<body[^>]*>)", r"\1" + overlay + _analista_live_panel(), html, count=1)
    return html


def _contexto_analista(db, eleccion, candidatos_dict: dict) -> dict:
    """Codex: resume el live dashboard en datos cerrados para el analista sin tokens."""
    if not eleccion:
        return {
            "ok": False,
            "motivo": "No hay eleccion activa",
            "total_votos": 0,
            "total_opiniones": 0,
        }

    eid = eleccion["id"]
    tipo_eleccion = eleccion["tipo"]
    datos_ventaja, datos_tendencia, total_votos = _datos_vivos(db, eid)
    fuente_datos = "live"
    usando_referencia = False
    if total_votos == 0:
        datos_ventaja, datos_tendencia, total_votos = _datos_dashboard_referencia(
            db, eid, nivel="estado", fuente="todos"
        )
        fuente_datos = "dashboard_referencia" if total_votos else "sin_datos"
        usando_referencia = total_votos > 0

    muestra_total = db.execute(
        "SELECT COUNT(*) c FROM muestra WHERE id_eleccion=? AND activo=1",
        (eid,),
    ).fetchone()["c"]
    if usando_referencia:
        centros_reportando = muestra_total
    else:
        centros_reportando = db.execute("""
            SELECT COUNT(DISTINCT v.codigo_centro) c
            FROM votos v
            JOIN muestra m ON m.codigo_centro = v.codigo_centro
            WHERE m.id_eleccion = ? AND v.valido = 1
        """, (eid,)).fetchone()["c"]

    puntos_nac = datos_tendencia.get("VENEZUELA") or []
    ventajas_nacionales = [
        round(float(p["gob"]) - float(p["opo"]), 1)
        for p in puntos_nac
    ]
    ventaja_actual = ventajas_nacionales[-1] if ventajas_nacionales else None
    hora_actual = puntos_nac[-1]["hora"] if puntos_nac else None

    if ventaja_actual is None:
        candidato_arriba = None
        candidato_abajo = None
    elif ventaja_actual >= 0:
        candidato_arriba = candidatos_dict.get("gobierno", "Gobierno")
        candidato_abajo = candidatos_dict.get("oposicion", "Oposicion")
    else:
        candidato_arriba = candidatos_dict.get("oposicion", "Oposicion")
        candidato_abajo = candidatos_dict.get("gobierno", "Gobierno")

    tendencias_por_estado = {k: v for k, v in datos_tendencia.items() if k != "VENEZUELA"}
    cortes_nacionales = len(puntos_nac)
    reglas = {
        "nacional": {"minimo_opiniones": 100, "minimo_cobertura_pct": 15, "minimo_cortes": 3},
        "regional": {"minimo_opiniones": 60, "minimo_cobertura_pct": 10, "minimo_cortes": 3},
        "municipal": {"minimo_opiniones": 30, "minimo_cobertura_pct": 10, "minimo_cortes": 3},
        "asamblea": {"minimo_opiniones": 60, "minimo_cobertura_pct": 10, "minimo_cortes": 3},
    }
    regla = reglas.get(tipo_eleccion, reglas["nacional"])
    estados_suf = {}
    if usando_referencia:
        for nombre, puntos in tendencias_por_estado.items():
            estados_suf[nombre.title()] = {
                "opiniones": None,
                "cobertura_pct": 100.0,
                "cortes": len(puntos),
                "datos_suficientes": len(puntos) >= regla["minimo_cortes"],
                "fuente_datos": fuente_datos,
                "nota": "Referencia de dashboard; no son opiniones SMS reales.",
            }
    else:
        rows_estado_suf = db.execute("""
            SELECT
                est.nombre AS estado,
                COUNT(v.id) AS opiniones,
                COUNT(DISTINCT v.codigo_centro) AS centros_reportando,
                COUNT(DISTINCT m.codigo_centro) AS centros_muestra,
                COUNT(DISTINCT v.turno) AS cortes
            FROM muestra m
            JOIN centros c ON c.codigo_cne = m.codigo_centro
            JOIN estados est ON est.id = c.id_estado
            LEFT JOIN votos v
              ON v.codigo_centro = m.codigo_centro
             AND v.valido = 1
            WHERE m.id_eleccion = ? AND m.activo = 1
            GROUP BY est.id
        """, (eid,)).fetchall()
        for r in rows_estado_suf:
            nombre = _norm_estado(r["estado"])
            centros_estado = int(r["centros_muestra"] or 0)
            cobertura_estado = round(100 * int(r["centros_reportando"] or 0) / centros_estado, 1) if centros_estado else 0
            opiniones_estado = int(r["opiniones"] or 0)
            cortes_estado = int(r["cortes"] or 0)
            estados_suf[nombre] = {
                "opiniones": opiniones_estado,
                "cobertura_pct": cobertura_estado,
                "cortes": cortes_estado,
                "datos_suficientes": (
                    opiniones_estado >= regla["minimo_opiniones"]
                    and cobertura_estado >= regla["minimo_cobertura_pct"]
                    and cortes_estado >= regla["minimo_cortes"]
                ),
            }
    datos_suficientes = (
        total_votos >= regla["minimo_opiniones"]
        and (round(100 * centros_reportando / muestra_total, 1) if muestra_total else 0) >= regla["minimo_cobertura_pct"]
        and cortes_nacionales >= regla["minimo_cortes"]
    )

    return {
        "ok": True,
        "eleccion": eleccion["nombre"],
        "tipo_eleccion": tipo_eleccion,
        "fuente_datos": fuente_datos,
        "nota_fuente": (
            "Datos de referencia del dashboard; no son opiniones recibidas por SMS en vivo."
            if usando_referencia else None
        ),
        "hora_actual": hora_actual,
        "total_votos": total_votos,
        "total_opiniones": total_votos,
        "centros_reportando": centros_reportando,
        "centros_muestra_total": muestra_total,
        "cobertura_pct": round(100 * centros_reportando / muestra_total, 1) if muestra_total else 0,
        "ventaja_actual": ventaja_actual,
        "ventajas_nacionales": ventajas_nacionales,
        "candidato_arriba": candidato_arriba,
        "candidato_abajo": candidato_abajo,
        "ventajas_por_estado": datos_ventaja,
        "tendencias_por_estado": tendencias_por_estado,
        "candidatos": candidatos_dict,
        "suficiencia": {
            **regla,
            "cortes": cortes_nacionales,
            "datos_suficientes": datos_suficientes,
            "estados": estados_suf,
        },
    }


def get_contexto_centro(centro_id: str | None = None) -> dict:
    db = get_db()
    try:
        eleccion = db.execute(
            "SELECT * FROM elecciones WHERE activa = 1 LIMIT 1"
        ).fetchone()
        if not eleccion:
            return {"ok": False, "motivo": "No hay eleccion activa"}

        eid = eleccion["id"]
        params: list = [eid]
        filtro_centro = ""
        if centro_id:
            filtro_centro = " AND v.codigo_centro = ?"
            params.append(centro_id)

        conteos = db.execute(f"""
            SELECT ca.id, ca.nombre, ca.bando, COUNT(v.id) opiniones
            FROM candidatos ca
            LEFT JOIN votos v
              ON v.id_candidato = ca.id
             AND v.valido = 1
             AND v.codigo_centro IN (
                SELECT codigo_centro FROM muestra WHERE id_eleccion = ?
             )
             {filtro_centro}
            WHERE ca.id_eleccion = ?
            GROUP BY ca.id
            ORDER BY ca.orden
        """, [eid, *([centro_id] if centro_id else []), eid]).fetchall()

        trend_params: list = [eid]
        filtro_turno = ""
        if centro_id:
            filtro_turno = " AND v.codigo_centro = ?"
            trend_params.append(centro_id)
        tendencias = db.execute(f"""
            SELECT v.turno, ca.nombre, ca.bando, COUNT(*) opiniones
            FROM votos v
            JOIN candidatos ca ON ca.id = v.id_candidato
            JOIN muestra m ON m.codigo_centro = v.codigo_centro
            WHERE m.id_eleccion = ? AND v.valido = 1 {filtro_turno}
              AND v.turno IN (
                SELECT turno FROM votos
                WHERE valido = 1
                GROUP BY turno
                ORDER BY turno DESC
                LIMIT 3
              )
            GROUP BY v.turno, ca.id
            ORDER BY v.turno DESC, ca.orden
        """, trend_params).fetchall()

        historico = None
        if centro_id:
            muestra = db.execute("""
                SELECT tipo_centro FROM muestra
                WHERE id_eleccion=? AND codigo_centro=?
                LIMIT 1
            """, (eid, centro_id)).fetchone()
            hist = db.execute("""
                SELECT eleccion_ref, pct_gobierno, pct_oposicion
                FROM resultados_historicos
                WHERE codigo_centro=?
                ORDER BY eleccion_ref DESC
                LIMIT 3
            """, (centro_id,)).fetchall()
            clasificacion = muestra["tipo_centro"] if muestra else None
            if not clasificacion and hist:
                diffs = [abs((r["pct_gobierno"] or 0) - (r["pct_oposicion"] or 0)) for r in hist]
                avg_diff = sum(diffs) / len(diffs)
                clasificacion = "swing" if avg_diff < 6 else "bastion"
            historico = {
                "clasificacion": clasificacion,
                "referencias": [dict(r) for r in hist],
            }

        conteos_list = [dict(r) for r in conteos]
        tendencias_list = [dict(r) for r in tendencias]
        total_opiniones = sum(int(r["opiniones"] or 0) for r in conteos_list)
        datos_suficientes = total_opiniones >= 100 and len({r["turno"] for r in tendencias_list}) >= 3

        return {
            "ok": True,
            "eleccion": eleccion["nombre"],
            "centro_id": centro_id,
            "conteos_de_opiniones_por_candidato": conteos_list,
            "ultimos_3_turnos": tendencias_list,
            "historico_centro": historico,
            "total_opiniones": total_opiniones,
            "datos_suficientes": datos_suficientes,
            "motivo_insuficiencia": None if datos_suficientes else "No hay suficientes opiniones y cortes comparables para analizar este ambito.",
        }
    finally:
        db.close()


@app.post("/chat")
async def chat(request: Request):
    payload = await request.json()
    question = (payload.get("question") or payload.get("pregunta") or "").strip()
    centro_id = payload.get("centro_id") or payload.get("codigo_centro")
    if not question:
        raise HTTPException(400, "Pregunta vacia")

    db = get_db()
    try:
        cfg_row = get_ai_config(db)
        provider = cfg_row["provider"]
        cfg = dict(cfg_row)
    finally:
        db.close()

    context = get_contexto_centro(centro_id)
    if not context.get("datos_suficientes"):
        return StreamingResponse(
            iter(["datos insuficientes para establecer tendencias"]),
            media_type="text/plain; charset=utf-8",
        )
    sys.path.insert(0, str(BASE_DIR))
    import agent

    def stream():
        try:
            yield from agent.ask_agent(question, context, provider, cfg)
        except Exception as exc:
            yield f"\n[error] {exc}"

    return StreamingResponse(stream(), media_type="text/plain; charset=utf-8")


@app.get("/config", response_class=HTMLResponse)
async def config_page(request: Request, msg: str = "", cat: str = "success"):
    db = get_db()
    try:
        ensure_config_table(db)
        rows = db.execute("SELECT * FROM config ORDER BY provider").fetchall()
        active = get_ai_config(db)
    finally:
        db.close()

    configs = {}
    for r in rows:
        item = dict(r)
        item.pop("api_key", None)
        configs[r["provider"]] = item
    return templates.TemplateResponse(request=request, name="config.html", context={
        "providers": AI_PROVIDER_DEFAULTS,
        "configs": configs,
        "active_provider": active["provider"] if active else "openai",
        "msg": msg,
        "cat": cat,
    })


@app.post("/config/guardar")
async def config_save(
    provider: str = Form(...),
    api_key: str = Form(""),
    model: str = Form(...),
    temperature: float = Form(0),
    max_tokens: int = Form(300),
):
    if provider not in AI_PROVIDER_DEFAULTS:
        raise HTTPException(400, "Proveedor no soportado")

    db = get_db()
    try:
        ensure_config_table(db)
        saved_key = api_key.strip() or None
        db.execute("UPDATE config SET active=0")
        db.execute("""
            INSERT INTO config (provider, api_key, model, temperature, max_tokens, active, updated_at)
            VALUES (?, ?, ?, ?, ?, 1, datetime('now'))
            ON CONFLICT(provider) DO UPDATE SET
                api_key=excluded.api_key,
                model=excluded.model,
                temperature=excluded.temperature,
                max_tokens=excluded.max_tokens,
                active=1,
                updated_at=datetime('now')
        """, (provider, saved_key, model.strip(), temperature, max_tokens))
        db.commit()
    finally:
        db.close()
    return RedirectResponse("/config?msg=Configuracion+guardada&cat=success", status_code=303)


@app.post("/config/test")
async def config_test(request: Request):
    db = get_db()
    try:
        cfg_row = get_ai_config(db)
        provider = cfg_row["provider"]
        cfg = dict(cfg_row)
    finally:
        db.close()

    sys.path.insert(0, str(BASE_DIR))
    import agent

    context = {
        "ok": True,
        "conteos_de_opiniones_por_candidato": [],
        "ultimos_3_turnos": [],
        "historico_centro": None,
        "nota": "Prueba fija de conexion sin datos electorales reales.",
        "datos_suficientes": True,
    }
    try:
        text = "".join(agent.ask_agent(
            "Responde solo si recibiste esta prueba de conexion.",
            context,
            provider,
            cfg,
        ))
        return JSONResponse({"ok": True, "provider": provider, "raw": text})
    except Exception as exc:
        return JSONResponse({"ok": False, "provider": provider, "raw": str(exc)}, status_code=502)


@app.get("/api/analista/contexto")
async def analista_contexto():
    db = get_db()
    try:
        eleccion = db.execute(
            "SELECT * FROM elecciones WHERE activa = 1 LIMIT 1"
        ).fetchone()
        candidatos_dict = _nombres_candidatos(db, eleccion["id"]) if eleccion else {}
        return JSONResponse(_contexto_analista(db, eleccion, candidatos_dict))
    finally:
        db.close()


@app.post("/api/analista/preguntar")
async def analista_preguntar(request: Request):
    payload = await request.json()
    pregunta = (payload.get("pregunta") or "").strip()
    db = get_db()
    try:
        eleccion = db.execute(
            "SELECT * FROM elecciones WHERE activa = 1 LIMIT 1"
        ).fetchone()
        candidatos_dict = _nombres_candidatos(db, eleccion["id"]) if eleccion else {}
        contexto = _contexto_analista(db, eleccion, candidatos_dict)

        sys.path.insert(0, str(BASE_DIR))
        import analista_ia
        return JSONResponse(analista_ia.analizar_contexto(contexto, pregunta))
    finally:
        db.close()


def _analista_live_panel() -> str:
    """Codex: panel persistente y aislado de las actualizaciones SSE."""
    return """
<!-- Codex: AI electoral analyst panel, deterministic and refresh-safe. -->
<style>
  #ai-analyst {
    position: fixed; right: 16px; bottom: 16px; z-index: 100000;
    width: min(390px, calc(100vw - 32px)); background: #fff; color: #1f2933;
    border: 1px solid rgba(0,0,0,.18); border-radius: 8px;
    box-shadow: 0 8px 28px rgba(0,0,0,.28); font-family: Arial, sans-serif;
  }
  #ai-analyst header {
    display: flex; align-items: center; justify-content: space-between;
    padding: 9px 11px; background: #1a3a5c; color: #fff;
    border-radius: 8px 8px 0 0; font-size: 13px; font-weight: 700;
  }
  #ai-analyst-log { max-height: 270px; overflow-y: auto; padding: 10px; font-size: 12px; }
  .ai-msg { margin-bottom: 8px; line-height: 1.35; }
  .ai-q { color: #1a3a5c; font-weight: 700; }
  .ai-a { background: #f4f7fb; border-left: 3px solid #1a3a5c; padding: 8px; border-radius: 4px; }
  #ai-analyst form { display: flex; gap: 6px; padding: 9px; border-top: 1px solid #e5e7eb; }
  #ai-analyst input { flex: 1; min-width: 0; padding: 7px; border: 1px solid #cbd5e1; border-radius: 5px; font-size: 12px; }
  #ai-analyst button { border: 0; border-radius: 5px; padding: 7px 9px; background: #1a3a5c; color: #fff; font-size: 12px; cursor: pointer; }
  #ai-analyst-clear { background: transparent !important; padding: 0 !important; color: #dbeafe !important; }
</style>
<section id="ai-analyst" aria-label="Analista electoral">
  <header>
    <span>AI Electoral Analyst</span>
    <button id="ai-analyst-clear" type="button">limpiar</button>
  </header>
  <div id="ai-analyst-log"></div>
  <form id="ai-analyst-form">
    <input id="ai-analyst-input" autocomplete="off" placeholder="Pregunta por la tendencia actual">
    <button type="submit">Analizar</button>
  </form>
</section>
<script>
(function() {
  var key = 'exitpoll.aiAnalyst.history';
  var draftKey = 'exitpoll.aiAnalyst.draft';
  var log = document.getElementById('ai-analyst-log');
  var input = document.getElementById('ai-analyst-input');
  var form = document.getElementById('ai-analyst-form');
  var clear = document.getElementById('ai-analyst-clear');

  function history() {
    try { return JSON.parse(localStorage.getItem(key) || '[]'); }
    catch(e) { return []; }
  }
  function save(items) { localStorage.setItem(key, JSON.stringify(items.slice(-8))); }
  function render() {
    var items = history();
    if (!items.length) {
      log.innerHTML = '<div class="ai-msg ai-a">Pregunta por ventaja, estabilidad o suficiencia de datos. No declaro ganadores.</div>';
      return;
    }
    log.innerHTML = items.map(function(item) {
      return '<div class="ai-msg"><div class="ai-q">' + escapeHtml(item.q) + '</div>' +
             '<div class="ai-a">' + escapeHtml(item.a) + '</div></div>';
    }).join('');
    log.scrollTop = log.scrollHeight;
  }
  function escapeHtml(s) {
    return String(s || '').replace(/[&<>"']/g, function(c) {
      return ({'&':'&amp;','<':'&lt;','>':'&gt;','"':'&quot;',"'":'&#39;'})[c];
    });
  }
  form.addEventListener('submit', function(e) {
    e.preventDefault();
    var q = (input.value || '').trim();
    if (!q) return;
    input.value = '';
    localStorage.removeItem(draftKey);
    fetch('/api/analista/preguntar', {
      method: 'POST',
      headers: {'Content-Type': 'application/json'},
      body: JSON.stringify({pregunta: q})
    }).then(function(r) { return r.json(); }).then(function(data) {
      var items = history();
      items.push({q: q, a: data.resumen || 'No hay lectura disponible.'});
      save(items);
      render();
    }).catch(function() {
      var items = history();
      items.push({q: q, a: 'No pude consultar el analista en este momento.'});
      save(items);
      render();
    });
  });
  input.value = localStorage.getItem(draftKey) || '';
  input.addEventListener('input', function() { localStorage.setItem(draftKey, input.value || ''); });
  clear.addEventListener('click', function() {
    localStorage.removeItem(key);
    localStorage.removeItem(draftKey);
    input.value = '';
    render();
  });
  render();
})();
</script>
"""


@app.get("/live", response_class=HTMLResponse)
async def live_dashboard(refresh: int = 5):
    """
    Dashboard en tiempo real.
    Abrir en el browser y correr simulador_showcase.py en otra terminal.
    El mapa y las tendencias se actualizan por SSE sin recargar la pagina.
    """
    db = get_db()
    try:
        eleccion = db.execute(
            "SELECT * FROM elecciones WHERE activa = 1 LIMIT 1"
        ).fetchone()

        if not eleccion:
            return HTMLResponse(_html_sin_datos("No hay elección activa", refresh))

        eid = eleccion["id"]
        datos_ventaja, datos_tendencia, total_votos = _datos_vivos(db, eid)
        fuente_datos = "live"
        if total_votos == 0:
            datos_ventaja, datos_tendencia, total_votos = _datos_dashboard_referencia(
                db, eid, nivel="estado", fuente="todos"
            )
            fuente_datos = "dashboard_referencia" if total_votos else "sin_datos"
        candidatos_dict = _nombres_candidatos(db, eid)
    finally:
        db.close()

    if total_votos == 0:
        return HTMLResponse(_html_sin_datos(
            f"Simulación no iniciada<br>"
            f"<small style='font-size:13px;opacity:.8'>{eleccion['nombre']}</small><br>"
            f"<small style='font-size:11px;opacity:.55'>"
            f"Corre: python backend/simulador_showcase.py --reset</small>",
            refresh,
        ))

    # ── Generar dashboard completo ────────────────────────────────
    sys.path.insert(0, str(BASE_DIR))
    import generador_dashboard
    import importlib
    importlib.reload(generador_dashboard)

    titulo = f"{eleccion['nombre']} — EN VIVO"
    subtitulo_fuente = (
        "Datos en vivo por SMS"
        if fuente_datos == "live"
        else "Datos de referencia del dashboard hasta recibir opiniones en vivo"
    )
    tmp = tempfile.mktemp(suffix=".html")
    try:
        generador_dashboard.generar_dashboard(
            datos_ventaja, datos_tendencia,
            nivel="estado", ruta_salida=tmp,
            titulo=titulo, candidatos=candidatos_dict,
        )
        with open(tmp, encoding="utf-8") as f:
            html = f.read()
    finally:
        if os.path.exists(tmp):
            os.unlink(tmp)

    # ── Inyectar barra de estado ───────────────────
    barra = (
        f'<style>:root{{--ep-top-offset:28px;}}</style>'
        f'<div style="position:fixed;top:0;left:0;right:0;z-index:99999;'
        f'background:#1B5E20;color:white;font-family:sans-serif;font-size:12px;'
        f'padding:5px 16px;display:flex;justify-content:space-between;align-items:center;'
        f'box-shadow:0 2px 6px rgba(0,0,0,.35);">'
        f'<span>&#x1F534;&nbsp; EN VIVO &nbsp;&#x2502;&nbsp; {eleccion["nombre"]}'
        f'&nbsp;&#x2502;&nbsp; <span id="ep-live-total">{total_votos:,} opiniones procesadas</span></span>'
        f'<span style="opacity:.7"><span id="ep-live-source">{subtitulo_fuente}</span> '
        f'&nbsp;&#x2502;&nbsp; SSE&nbsp;cada 60s</span>'
        f'</div>'
    )

    html = re.sub(r"(<body[^>]*>)", r"\1" + barra + _analista_live_panel(), html, count=1)

    return HTMLResponse(html)


# =============================================================
# TEST — carga de datos de demostración
# =============================================================

CSV_2024 = BASE_DIR / "resultados_cne2024.csv"
_CANDIDATOS_DEMO = [
    {"nombre": "Nicolás Maduro",   "partido": "PSUV", "bando": "gobierno",  "tipo": "unico", "orden": 1},
    {"nombre": "Edmundo González", "partido": "PUD",  "bando": "oposicion", "tipo": "unico", "orden": 2},
]


def _pct_2024() -> dict[str, tuple[float, float]]:
    pct: dict[str, tuple[float, float]] = {}
    with open(CSV_2024, encoding="utf-8-sig") as f:
        for row in csv.DictReader(f):
            pct[row["centro_cne_id"]] = (
                float(row["pct_gobierno"]) / 100,
                float(row["pct_oposicion"]) / 100,
            )
    return pct


def _asegurar_candidatos(db: sqlite3.Connection, id_eleccion: int) -> dict[str, int]:
    rows = db.execute(
        "SELECT id, bando FROM candidatos WHERE id_eleccion=?", (id_eleccion,)
    ).fetchall()
    if rows:
        return {r["bando"]: r["id"] for r in rows}
    for c in _CANDIDATOS_DEMO:
        db.execute(
            "INSERT INTO candidatos (id_eleccion, nombre, partido, bando, tipo, orden) "
            "VALUES (?,?,?,?,?,?)",
            (id_eleccion, c["nombre"], c["partido"], c["bando"], c["tipo"], c["orden"]),
        )
    db.commit()
    return {r["bando"]: r["id"] for r in db.execute(
        "SELECT id, bando FROM candidatos WHERE id_eleccion=?", (id_eleccion,)
    )}


def _insertar_votos(db: sqlite3.Connection, eleccion, centros, cands, pct, turnos_subset):
    for idx, hora_str in enumerate(turnos_subset):
        turno_num = idx + 1
        hora_iso  = f'{eleccion["fecha"]}T{hora_str}:00'
        for c in centros:
            cod  = c["codigo_centro"]
            p_g, p_o = pct.get(cod, (0.5, 0.5))
            p_tot = p_g + p_o
            tel  = f'+58414{c["id_muestra"]:07d}'
            lat  = c["lat"]  or (10.5  + random.uniform(-2, 2))
            lon  = c["lon"]  or (-66.9 + random.uniform(-3, 3))
            for _ in range(3):
                id_cand = (cands["gobierno"]
                           if random.random() < (p_g / p_tot)
                           else cands["oposicion"])
                cur = db.execute(
                    "INSERT INTO sms_raw (from_number, contenido, recibido_at, procesado) "
                    "VALUES (?,?,?,1)",
                    (tel, f"DEMO T{turno_num}", hora_iso),
                )
                db.execute(
                    "INSERT INTO votos (id_sms, codigo_centro, id_candidato, telefono, "
                    "hora, turno, lat, lon, distancia_m, valido) VALUES (?,?,?,?,?,?,?,?,?,1)",
                    (cur.lastrowid, cod, id_cand, tel,
                     hora_iso, turno_num, lat, lon, random.randint(30, 280)),
                )


def _encuestadores_demo(db: sqlite3.Connection, centros, id_eleccion: int):
    for c in centros:
        tel = f'+58414{c["id_muestra"]:07d}'
        db.execute(
            "INSERT OR IGNORE INTO encuestadores (telefono, nombre, codigo_centro, id_eleccion) "
            "VALUES (?,?,?,?)",
            (tel, f'Demo-{c["id_muestra"]}', c["codigo_centro"], id_eleccion),
        )


def _turnos(eleccion) -> list[str]:
    t      = datetime.strptime(eleccion["hora_apertura"], "%H:%M")
    cierre = datetime.strptime(eleccion["hora_cierre"],   "%H:%M")
    result = []
    while t <= cierre:
        result.append(t.strftime("%H:%M"))
        t += timedelta(minutes=20)
    return result


# ── Históricos ───────────────────────────────────────────────────────────────

def _datos_ventaja_historico_ref(conn, eleccion_ref: str) -> dict:
    rows = conn.execute("""
        SELECT e.nombre,
               SUM(rh.votos_gobierno)   AS gob,
               SUM(rh.votos_oposicion)  AS opo,
               SUM(rh.votos_validos)    AS val
        FROM resultados_historicos rh
        JOIN centros c ON rh.codigo_centro = c.codigo_cne
        JOIN estados e ON c.id_estado = e.id
        WHERE rh.eleccion_ref = ?
        GROUP BY e.id
    """, (eleccion_ref,)).fetchall()
    datos = {}
    for r in rows:
        if r["val"] and r["val"] > 0:
            datos[_norm_estado(r["nombre"])] = round(
                100 * r["gob"] / r["val"] - 100 * r["opo"] / r["val"], 1
            )
    return datos


def _historicos_unificados(conn: sqlite3.Connection) -> tuple[list[dict], dict[str, Any]]:
    """Construye la lista comun que consumen /historicos y debug-json."""
    diag: dict[str, Any] = {
        "est_rows": [],
        "rh_rows": [],
        "est_error": None,
        "rh_error": None,
    }

    def titulo_ref(ref: str) -> str:
        partes = ref.split("-", 1)
        if len(partes) == 2 and partes[0].isdigit():
            return f"{partes[1].replace('-', ' ').title()} {partes[0]}"
        return ref

    def orden_eleccion(row: dict) -> tuple[str, str]:
        ref = row.get("eleccion_ref") or ""
        year = ref.split("-", 1)[0] if ref[:4].isdigit() else "0000"
        return (row.get("fecha_eleccion") or year, ref)

    try:
        est_rows = conn.execute("""
            SELECT eleccion_ref,
                   COALESCE(MAX(CASE WHEN f='e' THEN nombre_eleccion END),
                            MAX(CASE WHEN f='o' THEN nombre_eleccion END)) AS nombre_eleccion,
                   COALESCE(MAX(CASE WHEN f='e' THEN fecha_eleccion END),
                            MAX(CASE WHEN f='o' THEN fecha_eleccion END))  AS fecha_eleccion,
                   MAX(CASE WHEN f='e' THEN pct_gov   END) AS e_gov,
                   MAX(CASE WHEN f='e' THEN pct_opos  END) AS e_opos,
                   MAX(CASE WHEN f='o' THEN pct_gov   END) AS o_gov,
                   MAX(CASE WHEN f='o' THEN pct_opos  END) AS o_opos,
                   MAX(CASE WHEN f='o' THEN pct_otros END) AS o_otros,
                   MAX(CASE WHEN f='e' THEN num_centros END) AS e_centros,
                   MAX(CASE WHEN f='o' THEN total_votos END) AS total_votos
            FROM (
                SELECT eleccion_ref, nombre_eleccion, fecha_eleccion, pct_gov, pct_opos,
                       NULL AS pct_otros, num_centros, NULL AS total_votos, 'e' AS f
                FROM historico_estudios WHERE ambito='NACIONAL'
                UNION ALL
                SELECT eleccion_ref, nombre_eleccion, fecha_eleccion, pct_gov, pct_opos,
                       pct_otros, NULL, total_votos, 'o'
                FROM historico_oficial WHERE ambito='NACIONAL'
            )
            GROUP BY eleccion_ref
            ORDER BY fecha_eleccion DESC, eleccion_ref DESC
        """).fetchall()
        diag["est_rows"] = [dict(r) for r in est_rows]
    except Exception as exc:
        est_rows = []
        diag["est_error"] = str(exc)

    est_refs = {r["eleccion_ref"] for r in est_rows}

    try:
        rh_rows = conn.execute("""
            SELECT eleccion_ref,
                   COUNT(*)                                                             AS num_centros,
                   SUM(votos_validos)                                                   AS total_votos,
                   ROUND(SUM(votos_gobierno)*100.0  / NULLIF(SUM(votos_validos),0), 1) AS o_gov,
                   ROUND(SUM(votos_oposicion)*100.0 / NULLIF(SUM(votos_validos),0), 1) AS o_opos
            FROM resultados_historicos
            GROUP BY eleccion_ref
            ORDER BY eleccion_ref DESC
        """).fetchall()
        diag["rh_rows"] = [dict(r) for r in rh_rows]
    except Exception as exc:
        rh_rows = []
        diag["rh_error"] = str(exc)

    elections: list[dict] = []
    for r in est_rows:
        d = dict(r)
        d["tipo"] = "con_estudio" if d["e_gov"] is not None else "oficial"
        d["nombre_eleccion"] = d.get("nombre_eleccion") or titulo_ref(d["eleccion_ref"])
        if d["e_gov"] is not None and d["o_gov"] is not None:
            d["delta"] = round(d["e_gov"] - d["o_gov"], 1)
            d["ganador_ok"] = (d["e_gov"] > d["e_opos"]) == (d["o_gov"] > d["o_opos"])
            # Indicador de riesgo de sesgo para la tarjeta (sin cálculo completo)
            delta_g = d["delta"]
            delta_o = round(d["e_opos"] - d["o_opos"], 1) if (d["e_opos"] is not None and d["o_opos"] is not None) else None
            errores_opuestos = (delta_g is not None and delta_o is not None and delta_g * delta_o < 0)
            if abs(delta_g) > 5:
                d["sesgo_nivel"] = "critico"
            elif abs(delta_g) > 2.5:
                d["sesgo_nivel"] = "alto" if errores_opuestos else "moderado"
            elif abs(delta_g) > 1:
                d["sesgo_nivel"] = "moderado" if errores_opuestos else "bajo"
            else:
                d["sesgo_nivel"] = "bajo"
        else:
            d["delta"] = None
            d["ganador_ok"] = None
            d["sesgo_nivel"] = None
        elections.append(d)

    for r in rh_rows:
        if r["eleccion_ref"] in est_refs:
            continue
        d = dict(r)
        d["tipo"] = "resultados"
        d["nombre_eleccion"] = titulo_ref(d["eleccion_ref"])
        d["fecha_eleccion"] = None
        d["o_otros"] = None
        d["e_gov"] = d["e_opos"] = d["delta"] = d["ganador_ok"] = None
        elections.append(d)

    # Inyectar tarjetas especiales para colecciones multi-estudio
    for ref, nombre, badge, url in [
        ("2008-gobernadores", "Elecciones Regionales 2008", "25 exit polls", "/historicos/estudios/2008-gobernadores"),
        ("2012-gobernadores", "Elecciones Regionales 2012", "23 exit polls", "/historicos/estudios/2012-gobernadores"),
        ("2013-municipales", "Elecciones Municipales 2013", "52 exit polls", "/historicos/estudios/2013-municipales"),
    ]:
        if any(e["eleccion_ref"] == ref for e in elections):
            for e in elections:
                if e["eleccion_ref"] != ref:
                    continue
                e["tipo"] = "coleccion"
                e["nombre_eleccion"] = nombre
                e["coleccion_url"] = url
                e["coleccion_badge"] = badge
                e["delta"] = None; e["ganador_ok"] = None; e["sesgo_nivel"] = None
                break

    elections.sort(key=orden_eleccion, reverse=True)
    diag["elections"] = elections
    return elections, diag


@app.get("/api/db-status", response_class=JSONResponse)
async def db_status():
    """Diagnóstico: conteo de filas por tabla relevante."""
    conn = get_db()
    tablas = [
        "resultados_historicos", "historico_estudios",
        "historico_oficial", "historico_estudios_turnos",
        "centros", "elecciones",
    ]
    out = {}
    for t in tablas:
        try:
            out[t] = conn.execute(f"SELECT COUNT(*) FROM {t}").fetchone()[0]
        except Exception as e:
            out[t] = f"ERROR: {e}"
    try:
        out["rh_refs"] = [r[0] for r in conn.execute(
            "SELECT DISTINCT eleccion_ref FROM resultados_historicos").fetchall()]
        out["he_refs"] = [r[0] for r in conn.execute(
            "SELECT DISTINCT eleccion_ref FROM historico_estudios").fetchall()]
    except Exception as e:
        out["refs_error"] = str(e)
    # Columnas reales de resultados_historicos (para detectar schema viejo)
    try:
        out["rh_columns"] = [r[1] for r in conn.execute(
            "PRAGMA table_info(resultados_historicos)").fetchall()]
    except Exception as e:
        out["rh_columns"] = f"ERROR: {e}"
    # Prueba la query de agregación que usa el route /historicos
    try:
        row = conn.execute("""
            SELECT COUNT(*) AS n,
                   ROUND(SUM(votos_gobierno)*100.0 / NULLIF(SUM(votos_validos),0), 1) AS o_gov,
                   ROUND(SUM(votos_oposicion)*100.0 / NULLIF(SUM(votos_validos),0), 1) AS o_opos
            FROM resultados_historicos
        """).fetchone()
        out["rh_agg_test"] = dict(row)
    except Exception as e:
        out["rh_agg_test"] = f"ERROR: {e}"
    conn.close()
    return out


@app.get("/historicos/debug-json", response_class=JSONResponse)
async def historicos_debug():
    """Devuelve el estado crudo del route /historicos para diagnóstico."""
    conn = get_db()
    _, out = _historicos_unificados(conn)
    out["deploy_ts"] = "2026-05-18T2018-presidencial"
    conn.close()
    return JSONResponse(out, headers={"Cache-Control": "no-store"})


@app.get("/historicos", response_class=HTMLResponse)
async def historicos_index(request: Request):
    conn = get_db()
    elections, diag = _historicos_unificados(conn)
    conn.close()
    response = templates.TemplateResponse(request=request, name="historicos.html", context={
        "elections": elections,
        "debug_counts": {
            "elections": len(elections),
            "est_rows": len(diag["est_rows"]),
            "rh_rows": len(diag["rh_rows"]),
        },
    })
    response.headers["Cache-Control"] = "no-store"
    return response


@app.get("/historicos/comparar", response_class=HTMLResponse)
async def historicos_comparar(request: Request, a: str = "", b: str = ""):
    conn = get_db()
    todos = conn.execute(
        "SELECT DISTINCT eleccion_ref FROM resultados_historicos ORDER BY eleccion_ref DESC"
    ).fetchall()

    comparacion = []
    if a and b:
        rows = conn.execute("""
            SELECT
                rh.eleccion_ref,
                e.nombre                                                            AS estado,
                ROUND(SUM(rh.votos_gobierno)*100.0  / NULLIF(SUM(rh.votos_validos),0),1) AS pct_gov,
                ROUND(SUM(rh.votos_oposicion)*100.0 / NULLIF(SUM(rh.votos_validos),0),1) AS pct_opos
            FROM resultados_historicos rh
            JOIN centros c ON rh.codigo_centro = c.codigo_cne
            JOIN estados e ON c.id_estado = e.id
            WHERE rh.eleccion_ref IN (?, ?)
            GROUP BY rh.eleccion_ref, e.id
            ORDER BY e.nombre
        """, (a, b)).fetchall()

        pivot: dict = {}
        for r in rows:
            pivot.setdefault(r["estado"], {})[r["eleccion_ref"]] = {
                "pct_gov": r["pct_gov"], "pct_opos": r["pct_opos"]
            }

        for estado, data in sorted(pivot.items()):
            da = data.get(a, {})
            db2 = data.get(b, {})
            gov_a, gov_b = da.get("pct_gov"), db2.get("pct_gov")
            swing = round(gov_b - gov_a, 1) if gov_a is not None and gov_b is not None else None
            comparacion.append({
                "estado": estado,
                "a_gov": gov_a, "a_opos": da.get("pct_opos"),
                "b_gov": gov_b, "b_opos": db2.get("pct_opos"),
                "swing": swing,
            })

    conn.close()
    return templates.TemplateResponse(request=request, name="historico_comparar.html", context={
        "todos": todos, "a": a, "b": b, "comparacion": comparacion
    })


# ── Historicos: Estudios ──────────────────────────────────────────────────────

_ESTADOS_FALLBACK = [
    ("01", "Distrito Capital"), ("02", "Amazonas"), ("03", "Anzoátegui"),
    ("04", "Apure"), ("05", "Aragua"), ("06", "Barinas"), ("07", "Bolívar"),
    ("08", "Carabobo"), ("09", "Cojedes"), ("10", "Delta Amacuro"),
    ("11", "Falcón"), ("12", "Guárico"), ("13", "La Guaira"),
    ("14", "Lara"), ("15", "Mérida"), ("16", "Miranda"), ("17", "Monagas"),
    ("18", "Nueva Esparta"), ("19", "Portuguesa"), ("20", "Sucre"),
    ("21", "Táchira"), ("22", "Trujillo"), ("23", "Yaracuy"), ("24", "Zulia"),
]


def _estados_para_form(conn) -> list[tuple[str, str]]:
    rows = conn.execute("SELECT codigo_cne, nombre FROM estados ORDER BY nombre").fetchall()
    if rows:
        return [(r["codigo_cne"], r["nombre"]) for r in rows]
    return _ESTADOS_FALLBACK


def _estudios_pivot(conn, ref: str) -> list[dict]:
    """LEFT+RIGHT JOIN emulado via UNION para SQLite (no soporta FULL OUTER JOIN)."""
    rows = conn.execute("""
        SELECT ambito, nombre,
               MAX(CASE WHEN fuente='e' THEN pct_gov   END) AS e_gov,
               MAX(CASE WHEN fuente='e' THEN pct_opos  END) AS e_opos,
               MAX(CASE WHEN fuente='e' THEN pct_otros END) AS e_otros,
               MAX(CASE WHEN fuente='e' THEN num_centros END) AS e_centros,
               MAX(CASE WHEN fuente='o' THEN pct_gov   END) AS o_gov,
               MAX(CASE WHEN fuente='o' THEN pct_opos  END) AS o_opos,
               MAX(CASE WHEN fuente='o' THEN pct_otros END) AS o_otros,
               MAX(CASE WHEN fuente='o' THEN total_votos END) AS o_votos
        FROM (
            SELECT ambito, nombre, pct_gov, pct_opos, pct_otros,
                   num_centros, NULL AS total_votos, 'e' AS fuente
            FROM historico_estudios WHERE eleccion_ref = ?
            UNION ALL
            SELECT ambito, nombre, pct_gov, pct_opos, pct_otros,
                   NULL, total_votos, 'o'
            FROM historico_oficial WHERE eleccion_ref = ?
        )
        GROUP BY ambito
        ORDER BY CASE WHEN ambito='NACIONAL' THEN 0 ELSE 1 END, nombre
    """, (ref, ref)).fetchall()

    result = []
    for r in rows:
        delta = None
        if r["e_gov"] is not None and r["o_gov"] is not None:
            delta = round(r["e_gov"] - r["o_gov"], 1)
        result.append({
            "ambito": r["ambito"], "nombre": r["nombre"],
            "e_gov": r["e_gov"], "e_opos": r["e_opos"], "e_otros": r["e_otros"],
            "e_centros": r["e_centros"],
            "o_gov": r["o_gov"], "o_opos": r["o_opos"], "o_otros": r["o_otros"],
            "o_votos": r["o_votos"],
            "delta_gov": delta,
            "error_abs": abs(delta) if delta is not None else None,
        })
    return result


@app.get("/historicos/estudios", response_class=HTMLResponse)
async def historico_estudios_index(request: Request):
    return RedirectResponse("/historicos", status_code=301)


@app.get("/historicos/estudios/nuevo", response_class=HTMLResponse)
async def historico_estudio_nuevo_get(request: Request, ref: str = ""):
    raise HTTPException(404, "Los estudios historicos son de solo lectura")
    conn = get_db()
    estados = _estados_para_form(conn)
    conn.close()
    return templates.TemplateResponse(request=request, name="historico_estudio_editar.html", context={
        "estados": estados, "estudio": {}, "oficial": {}, "turnos": {},
        "meta": {"ref": ref}, "edit_ref": ""
    })


@app.get("/historicos/estudios/{ref}/editar", response_class=HTMLResponse)
async def historico_estudio_editar_get(request: Request, ref: str):
    raise HTTPException(404, "Los estudios historicos son de solo lectura")
    conn = get_db()
    estados = _estados_para_form(conn)
    estudio = {r["ambito"]: dict(r) for r in
               conn.execute("SELECT * FROM historico_estudios WHERE eleccion_ref=?", (ref,)).fetchall()}
    oficial = {r["ambito"]: dict(r) for r in
               conn.execute("SELECT * FROM historico_oficial WHERE eleccion_ref=?", (ref,)).fetchall()}
    turnos_list = conn.execute(
        "SELECT * FROM historico_estudios_turnos WHERE eleccion_ref=? ORDER BY turno",
        (ref,)
    ).fetchall()
    turnos = {r["turno"]: dict(r) for r in turnos_list}
    nac_e = estudio.get("NACIONAL", {})
    nac_o = oficial.get("NACIONAL", {})
    meta = {
        "ref": ref,
        "nombre_eleccion": nac_e.get("nombre_eleccion") or nac_o.get("nombre_eleccion") or "",
        "fecha_eleccion":  nac_e.get("fecha_eleccion")  or nac_o.get("fecha_eleccion")  or "",
    }
    conn.close()
    return templates.TemplateResponse(request=request, name="historico_estudio_editar.html", context={
        "estados": estados, "estudio": estudio, "oficial": oficial,
        "turnos": turnos, "meta": meta, "edit_ref": ref
    })


@app.post("/historicos/estudios/guardar", response_class=RedirectResponse)
async def historico_estudio_guardar(request: Request):
    raise HTTPException(404, "Los estudios historicos son de solo lectura")
    form = await request.form()
    ref             = (form.get("eleccion_ref") or "").strip()
    nombre_eleccion = (form.get("nombre_eleccion") or "").strip()
    fecha_eleccion  = (form.get("fecha_eleccion") or "").strip()
    if not ref:
        raise HTTPException(400, "Referencia requerida")

    conn = get_db()
    estados = _estados_para_form(conn)
    ambitos = [("NACIONAL", "Nacional")] + list(estados)

    def _f(key):
        v = (form.get(key) or "").strip()
        try:
            return float(v) if v else None
        except ValueError:
            return None

    def _i(key):
        v = (form.get(key) or "").strip()
        try:
            return int(float(v)) if v else 0
        except ValueError:
            return 0

    # Parse & validate all ambito rows
    parsed_ambitos = []
    errores = []
    for ambito, nombre_amb in ambitos:
        e = {"gov": _f(f"e_{ambito}_gov"), "opos": _f(f"e_{ambito}_opos"),
             "otros": _f(f"e_{ambito}_otros") or 0,
             "centros": _i(f"e_{ambito}_centros"),
             "fuente": (form.get(f"e_{ambito}_fuente") or "").strip() or None}
        o = {"gov": _f(f"o_{ambito}_gov"), "opos": _f(f"o_{ambito}_opos"),
             "otros": _f(f"o_{ambito}_otros") or 0,
             "votos": _i(f"o_{ambito}_votos")}
        if e["gov"] is not None or e["opos"] is not None:
            total = (e["gov"] or 0) + (e["opos"] or 0) + e["otros"]
            if total > 0 and not (99.5 <= total <= 100.5):
                errores.append(f"{nombre_amb} (estudio): suma {total:.1f}% ≠ 100%")
        if o["gov"] is not None or o["opos"] is not None:
            total = (o["gov"] or 0) + (o["opos"] or 0) + o["otros"]
            if total > 0 and not (99.5 <= total <= 100.5):
                errores.append(f"{nombre_amb} (oficial): suma {total:.1f}% ≠ 100%")
        parsed_ambitos.append((ambito, nombre_amb, e, o))

    # Parse & validate turnos
    parsed_turnos = []
    for t in range(1, 13):
        tgov  = _f(f"t_{t}_gov")
        topos = _f(f"t_{t}_opos")
        totros = _f(f"t_{t}_otros") or 0
        if tgov is not None or topos is not None:
            total = (tgov or 0) + (topos or 0) + totros
            if total > 0 and not (99.5 <= total <= 100.5):
                errores.append(f"Turno {t}: suma {total:.1f}% ≠ 100%")
            parsed_turnos.append({
                "turno": t,
                "hora_label": (form.get(f"t_{t}_hora") or "").strip() or None,
                "gov": tgov or 0, "opos": topos or 0, "otros": totros,
                "centros": _i(f"t_{t}_centros"),
            })

    if errores:
        conn.close()
        raise HTTPException(400, "; ".join(errores))

    # Save ambito rows
    for ambito, nombre_amb, e, o in parsed_ambitos:
        is_nac = ambito == "NACIONAL"
        ne = nombre_eleccion if is_nac else None
        fe = fecha_eleccion if is_nac else None
        if e["gov"] is not None or e["opos"] is not None:
            conn.execute("""
                INSERT INTO historico_estudios
                    (eleccion_ref, ambito, nombre, nombre_eleccion, fecha_eleccion,
                     pct_gov, pct_opos, pct_otros, num_centros, fuente, updated_at)
                VALUES (?,?,?,?,?,?,?,?,?,?,datetime('now'))
                ON CONFLICT(eleccion_ref, ambito) DO UPDATE SET
                    nombre=excluded.nombre,
                    nombre_eleccion=COALESCE(excluded.nombre_eleccion, nombre_eleccion),
                    fecha_eleccion=COALESCE(excluded.fecha_eleccion, fecha_eleccion),
                    pct_gov=excluded.pct_gov, pct_opos=excluded.pct_opos,
                    pct_otros=excluded.pct_otros, num_centros=excluded.num_centros,
                    fuente=COALESCE(excluded.fuente, fuente), updated_at=excluded.updated_at
            """, (ref, ambito, nombre_amb, ne, fe,
                  e["gov"] or 0, e["opos"] or 0, e["otros"], e["centros"], e["fuente"]))
        if o["gov"] is not None or o["opos"] is not None:
            conn.execute("""
                INSERT INTO historico_oficial
                    (eleccion_ref, ambito, nombre, nombre_eleccion, fecha_eleccion,
                     pct_gov, pct_opos, pct_otros, total_votos, updated_at)
                VALUES (?,?,?,?,?,?,?,?,?,datetime('now'))
                ON CONFLICT(eleccion_ref, ambito) DO UPDATE SET
                    nombre=excluded.nombre,
                    nombre_eleccion=COALESCE(excluded.nombre_eleccion, nombre_eleccion),
                    fecha_eleccion=COALESCE(excluded.fecha_eleccion, fecha_eleccion),
                    pct_gov=excluded.pct_gov, pct_opos=excluded.pct_opos,
                    pct_otros=excluded.pct_otros, total_votos=excluded.total_votos,
                    updated_at=excluded.updated_at
            """, (ref, ambito, nombre_amb, ne, fe,
                  o["gov"] or 0, o["opos"] or 0, o["otros"], o["votos"]))

    # Save turnos
    for pt in parsed_turnos:
        conn.execute("""
            INSERT INTO historico_estudios_turnos
                (eleccion_ref, ambito, turno, hora_label, pct_gov, pct_opos, pct_otros, num_centros, updated_at)
            VALUES (?,?,?,?,?,?,?,?,datetime('now'))
            ON CONFLICT(eleccion_ref, ambito, turno) DO UPDATE SET
                hora_label=excluded.hora_label,
                pct_gov=excluded.pct_gov, pct_opos=excluded.pct_opos,
                pct_otros=excluded.pct_otros, num_centros=excluded.num_centros,
                updated_at=excluded.updated_at
        """, (ref, 'NACIONAL', pt["turno"], pt["hora_label"], pt["gov"], pt["opos"], pt["otros"], pt["centros"]))

    conn.commit()
    conn.close()
    return RedirectResponse(f"/historicos/estudios/{ref}", status_code=303)


@app.get("/historicos/estudios/2008-gobernadores", response_class=HTMLResponse)
async def gobernadores_2008_collection(request: Request):
    """Vista colección: mapa + grilla de los 25 exit polls regionales 2008."""
    conn = get_db()
    import json as _json, math as _math

    nac_row = conn.execute(
        "SELECT notas FROM historico_estudios WHERE eleccion_ref='2008-gobernadores' AND ambito='NACIONAL'"
    ).fetchone()
    if not nac_row:
        conn.close()
        raise HTTPException(404, "Datos de 2008-gobernadores no encontrados. Ejecute import_2008.py primero.")
    coleccion_meta = _json.loads(nac_row["notas"] or "{}")

    filas_e = {r["ambito"]: dict(r) for r in conn.execute(
        "SELECT ambito, nombre, pct_gov, pct_opos, pct_otros, num_centros, notas "
        "FROM historico_estudios WHERE eleccion_ref='2008-gobernadores' AND ambito!='NACIONAL'"
    ).fetchall()}
    filas_o = {r["ambito"]: dict(r) for r in conn.execute(
        "SELECT ambito, pct_gov, pct_opos, pct_otros "
        "FROM historico_oficial WHERE eleccion_ref='2008-gobernadores'"
    ).fetchall()}
    conn.close()

    ICC_REF = 0.04
    estados = []
    for slug, fe in filas_e.items():
        fo = filas_o.get(slug, {})
        notas = _json.loads(fe.get("notas") or "{}")
        e_gov  = fe.get("pct_gov");  e_opos = fe.get("pct_opos")
        o_gov  = fo.get("pct_gov");  o_opos = fo.get("pct_opos")
        n      = notas.get("n_respondentes", 0)
        k      = fe.get("num_centros", 0)
        deff   = round(1 + (n / k - 1) * ICC_REF, 2) if k > 0 else None
        moe    = round(1.96 * _math.sqrt((deff or 1) * 0.25 / n) * 100, 2) if n > 0 else None

        win_e  = e_gov is not None and e_opos is not None and e_gov > e_opos
        win_o  = o_gov is not None and o_opos is not None and o_gov > o_opos
        delta_g = round(e_gov - o_gov, 2) if e_gov is not None and o_gov else None
        delta_o = round(e_opos - o_opos, 2) if e_opos is not None and o_opos else None
        ganador_ok = (win_e == win_o) if o_gov else None
        lara_ambiguo = bool(notas.get("lara_nota"))

        estados.append({
            "slug": slug,
            "nombre": fe.get("nombre", slug),
            "region": notas.get("region", ""),
            "tipo_cargo": notas.get("tipo_cargo", "Gobernación"),
            "cand_gov": notas.get("cand_gov_nombre") or notas.get("candidato_gov"),
            "cand_opos": notas.get("cand_opos_nombre") or notas.get("candidato_opos"),
            "e_gov": e_gov, "e_opos": e_opos,
            "o_gov": o_gov, "o_opos": o_opos,
            "n_centros": k, "n_respondentes": n,
            "max_turno": 0,
            "deff": deff, "moe_pp": moe,
            "win_estudio": win_e, "win_oficial": win_o,
            "ganador_ok": ganador_ok, "lara_ambiguo": lara_ambiguo,
            "delta_gov": delta_g, "delta_opos": delta_o,
        })

    REGION_ORDER = ["Capital", "Central", "Centro Occidental",
                    "Andina", "Guayana y Llanos", "Nororiental e Insular", "Zuliana"]
    from collections import defaultdict
    por_region: dict = defaultdict(list)
    for e in estados:
        por_region[e["region"]].append(e)
    regiones = [(r, sorted(por_region[r], key=lambda x: x["nombre"])) for r in REGION_ORDER if r in por_region]
    correctos = sum(1 for e in estados if e["ganador_ok"] is True)
    con_oficial = sum(1 for e in estados if e["ganador_ok"] is not None)

    return templates.TemplateResponse(request=request,
        name="gobernadores_2008.html", context={
            "meta": coleccion_meta,
            "estados": estados,
            "regiones": regiones,
            "total_centros": coleccion_meta.get("n_centros_total", 0),
            "total_resp": coleccion_meta.get("n_respondentes_total", 0),
            "n_estudios": len(estados),
            "correctos": correctos,
            "con_oficial": con_oficial,
        })


@app.get("/historicos/estudios/2008-gobernadores/{estado_slug}", response_class=HTMLResponse)
async def gobernadores_2008_detalle(request: Request, estado_slug: str):
    """Detalle individual: turno-chart + DEFF + análisis acertividad por estado."""
    conn = get_db()
    import json as _json, math as _math

    fe = conn.execute(
        "SELECT * FROM historico_estudios WHERE eleccion_ref='2008-gobernadores' AND ambito=?",
        (estado_slug,)
    ).fetchone()
    if not fe:
        conn.close()
        raise HTTPException(404, f"Estado '{estado_slug}' no encontrado en 2008-gobernadores")
    fe = dict(fe)
    fo = conn.execute(
        "SELECT * FROM historico_oficial WHERE eleccion_ref='2008-gobernadores' AND ambito=?",
        (estado_slug,)
    ).fetchone()
    fo = dict(fo) if fo else {}

    turnos_raw = conn.execute(
        "SELECT turno, pct_gov, pct_opos, pct_otros, num_centros "
        "FROM historico_estudios_turnos "
        "WHERE eleccion_ref='2008-gobernadores' AND ambito=? ORDER BY turno",
        (estado_slug,)
    ).fetchall()
    conn.close()

    notas = _json.loads(fe.get("notas") or "{}")
    turnos = [dict(r) for r in turnos_raw]
    ICC_REF = 0.04

    e_gov  = fe.get("pct_gov");  e_opos = fe.get("pct_opos");  e_otros = fe.get("pct_otros")
    o_gov  = fo.get("pct_gov");  o_opos = fo.get("pct_opos");  o_otros = fo.get("pct_otros")
    n      = notas.get("n_respondentes", 0)
    k      = fe.get("num_centros", 0)
    deff   = round(1 + (n / k - 1) * ICC_REF, 2) if k > 0 else None
    moe_srs = round(1.96 * _math.sqrt(0.25 / n) * 100, 2) if n > 0 else None
    moe_adj = round(1.96 * _math.sqrt((deff or 1) * 0.25 / n) * 100, 2) if n > 0 else None

    analisis = {}
    if e_gov is not None and o_gov is not None and o_gov > 0:
        delta_g = round(e_gov - o_gov, 2)
        delta_o = round(e_opos - o_opos, 2) if e_opos is not None and o_opos else None
        e_brecha = round(e_gov - (e_opos or 0), 2)
        o_brecha = round(o_gov - (o_opos or 0), 2)
        error_abs = abs(delta_g)
        magnitud = ("leve" if error_abs <= 1.0 else "moderado" if error_abs <= 2.5
                    else "severo" if error_abs <= 5.0 else "critico")
        win_e = e_gov > (e_opos or 0)
        win_o = o_gov > (o_opos or 0)
        mae_partes = [error_abs]
        if delta_o is not None: mae_partes.append(abs(delta_o))
        if e_otros is not None and o_otros is not None and o_otros > 0:
            mae_partes.append(abs(round(e_otros - o_otros, 2)))
        analisis = {
            "delta_gov": delta_g, "delta_opos": delta_o,
            "e_brecha": e_brecha, "o_brecha": o_brecha,
            "delta_brecha": round(e_brecha - o_brecha, 2),
            "error_abs": error_abs, "magnitud": magnitud,
            "mae_mosteller": round(sum(mae_partes) / len(mae_partes), 2),
            "mae_n_opciones": len(mae_partes),
            "ganador_estudio": "gobierno" if win_e else "oposicion",
            "ganador_oficial": "gobierno" if win_o else "oposicion",
            "acierto_ganador": win_e == win_o,
            "errores_opuestos": (delta_g * delta_o < 0) if delta_o else False,
        }

    import json as json_mod
    return templates.TemplateResponse(request=request,
        name="gobernador_2008_detalle.html", context={
            "slug": estado_slug,
            "nombre": fe.get("nombre", estado_slug),
            "notas": notas,
            "fe": fe, "fo": fo,
            "e_gov": e_gov, "e_opos": e_opos, "e_otros": e_otros,
            "o_gov": o_gov, "o_opos": o_opos, "o_otros": o_otros,
            "n_respondentes": n, "n_centros": k,
            "deff": deff, "moe_srs": moe_srs, "moe_adj": moe_adj,
            "turnos_json": json_mod.dumps(turnos),
            "turnos": turnos,
            "analisis": analisis,
            "lara_ambiguo": bool(notas.get("lara_nota")),
        })


@app.get("/historicos/estudios/2012-gobernadores", response_class=HTMLResponse)
async def gobernadores_2012_collection(request: Request):
    """Vista coleccion: mapa + grilla de los 23 exit polls regionales 2012."""
    conn = get_db()
    import json as _json, math as _math

    nac_row = conn.execute(
        "SELECT notas FROM historico_estudios WHERE eleccion_ref='2012-gobernadores' AND ambito='NACIONAL'"
    ).fetchone()
    if not nac_row:
        conn.close()
        raise HTTPException(404, "Datos de 2012-gobernadores no encontrados. Ejecute import_2012_gobernadores.py primero.")
    coleccion_meta = _json.loads(nac_row["notas"] or "{}")

    filas_e = {r["ambito"]: dict(r) for r in conn.execute(
        "SELECT ambito, nombre, pct_gov, pct_opos, pct_otros, num_centros, notas "
        "FROM historico_estudios WHERE eleccion_ref='2012-gobernadores' AND ambito!='NACIONAL'"
    ).fetchall()}
    filas_o = {r["ambito"]: dict(r) for r in conn.execute(
        "SELECT ambito, pct_gov, pct_opos, pct_otros "
        "FROM historico_oficial WHERE eleccion_ref='2012-gobernadores'"
    ).fetchall()}
    conn.close()

    ICC_REF = 0.04
    estados = []
    for slug, fe in filas_e.items():
        fo = filas_o.get(slug, {})
        notas = _json.loads(fe.get("notas") or "{}")
        e_gov  = fe.get("pct_gov");  e_opos = fe.get("pct_opos")
        o_gov  = fo.get("pct_gov");  o_opos = fo.get("pct_opos")
        n      = notas.get("n_respondentes", 0)
        k      = fe.get("num_centros", 0)
        deff   = round(1 + (n / k - 1) * ICC_REF, 2) if k > 0 else None
        moe    = round(1.96 * _math.sqrt((deff or 1) * 0.25 / n) * 100, 2) if n > 0 else None

        win_e  = e_gov is not None and e_opos is not None and e_gov > e_opos
        win_o  = o_gov is not None and o_opos is not None and o_gov > o_opos
        delta_g = round(e_gov - o_gov, 2) if e_gov is not None and o_gov else None
        delta_o = round(e_opos - o_opos, 2) if e_opos is not None and o_opos else None
        ganador_ok = (win_e == win_o) if o_gov else None
        lara_ambiguo = bool(notas.get("lara_nota"))

        estados.append({
            "slug": slug,
            "nombre": fe.get("nombre", slug),
            "region": notas.get("region", ""),
            "tipo_cargo": notas.get("tipo_cargo", "Gobernacion"),
            "cand_gov": notas.get("cand_gov_nombre") or notas.get("candidato_gov"),
            "cand_opos": notas.get("cand_opos_nombre") or notas.get("candidato_opos"),
            "e_gov": e_gov, "e_opos": e_opos,
            "o_gov": o_gov, "o_opos": o_opos,
            "n_centros": k, "n_respondentes": n,
            "max_turno": 0,
            "deff": deff, "moe_pp": moe,
            "win_estudio": win_e, "win_oficial": win_o,
            "ganador_ok": ganador_ok, "lara_ambiguo": lara_ambiguo,
            "delta_gov": delta_g, "delta_opos": delta_o,
        })

    REGION_ORDER = ["Capital", "Central", "Centro Occidental",
                    "Andina", "Guayana y Llanos", "Nororiental e Insular", "Zuliana"]
    from collections import defaultdict
    por_region: dict = defaultdict(list)
    for e in estados:
        por_region[e["region"]].append(e)
    regiones = [(r, sorted(por_region[r], key=lambda x: x["nombre"])) for r in REGION_ORDER if r in por_region]
    correctos = sum(1 for e in estados if e["ganador_ok"] is True)
    con_oficial = sum(1 for e in estados if e["ganador_ok"] is not None)

    return templates.TemplateResponse(request=request,
        name="gobernadores_2012.html", context={
            "meta": coleccion_meta,
            "estados": estados,
            "regiones": regiones,
            "total_centros": coleccion_meta.get("n_centros_total", 0),
            "total_resp": coleccion_meta.get("n_respondentes_total", 0),
            "n_estudios": len(estados),
            "correctos": correctos,
            "con_oficial": con_oficial,
        })


@app.get("/historicos/estudios/2012-gobernadores/{estado_slug}", response_class=HTMLResponse)
async def gobernadores_2012_detalle(request: Request, estado_slug: str):
    """Detalle individual: turno-chart + DEFF + analisis acertividad por estado."""
    conn = get_db()
    import json as _json, math as _math

    fe = conn.execute(
        "SELECT * FROM historico_estudios WHERE eleccion_ref='2012-gobernadores' AND ambito=?",
        (estado_slug,)
    ).fetchone()
    if not fe:
        conn.close()
        raise HTTPException(404, f"Estado '{estado_slug}' no encontrado en 2012-gobernadores")
    fe = dict(fe)
    fo = conn.execute(
        "SELECT * FROM historico_oficial WHERE eleccion_ref='2012-gobernadores' AND ambito=?",
        (estado_slug,)
    ).fetchone()
    fo = dict(fo) if fo else {}

    turnos_raw = conn.execute(
        "SELECT turno, pct_gov, pct_opos, pct_otros, num_centros "
        "FROM historico_estudios_turnos "
        "WHERE eleccion_ref='2012-gobernadores' AND ambito=? ORDER BY turno",
        (estado_slug,)
    ).fetchall()
    conn.close()

    notas = _json.loads(fe.get("notas") or "{}")
    turnos = [dict(r) for r in turnos_raw]
    ICC_REF = 0.04

    e_gov  = fe.get("pct_gov");  e_opos = fe.get("pct_opos");  e_otros = fe.get("pct_otros")
    o_gov  = fo.get("pct_gov");  o_opos = fo.get("pct_opos");  o_otros = fo.get("pct_otros")
    n      = notas.get("n_respondentes", 0)
    k      = fe.get("num_centros", 0)
    deff   = round(1 + (n / k - 1) * ICC_REF, 2) if k > 0 else None
    moe_srs = round(1.96 * _math.sqrt(0.25 / n) * 100, 2) if n > 0 else None
    moe_adj = round(1.96 * _math.sqrt((deff or 1) * 0.25 / n) * 100, 2) if n > 0 else None

    analisis = {}
    if e_gov is not None and o_gov is not None and o_gov > 0:
        delta_g = round(e_gov - o_gov, 2)
        delta_o = round(e_opos - o_opos, 2) if e_opos is not None and o_opos else None
        e_brecha = round(e_gov - (e_opos or 0), 2)
        o_brecha = round(o_gov - (o_opos or 0), 2)
        error_abs = abs(delta_g)
        magnitud = ("leve" if error_abs <= 1.0 else "moderado" if error_abs <= 2.5
                    else "severo" if error_abs <= 5.0 else "critico")
        win_e = e_gov > (e_opos or 0)
        win_o = o_gov > (o_opos or 0)
        mae_partes = [error_abs]
        if delta_o is not None: mae_partes.append(abs(delta_o))
        if e_otros is not None and o_otros is not None and o_otros > 0:
            mae_partes.append(abs(round(e_otros - o_otros, 2)))
        analisis = {
            "delta_gov": delta_g, "delta_opos": delta_o,
            "e_brecha": e_brecha, "o_brecha": o_brecha,
            "delta_brecha": round(e_brecha - o_brecha, 2),
            "error_abs": error_abs, "magnitud": magnitud,
            "mae_mosteller": round(sum(mae_partes) / len(mae_partes), 2),
            "mae_n_opciones": len(mae_partes),
            "ganador_estudio": "gobierno" if win_e else "oposicion",
            "ganador_oficial": "gobierno" if win_o else "oposicion",
            "acierto_ganador": win_e == win_o,
            "errores_opuestos": (delta_g * delta_o < 0) if delta_o else False,
        }

    import json as json_mod
    return templates.TemplateResponse(request=request,
        name="gobernador_2012_detalle.html", context={
            "slug": estado_slug,
            "nombre": fe.get("nombre", estado_slug),
            "notas": notas,
            "fe": fe, "fo": fo,
            "e_gov": e_gov, "e_opos": e_opos, "e_otros": e_otros,
            "o_gov": o_gov, "o_opos": o_opos, "o_otros": o_otros,
            "n_respondentes": n, "n_centros": k,
            "deff": deff, "moe_srs": moe_srs, "moe_adj": moe_adj,
            "turnos_json": json_mod.dumps(turnos),
            "turnos": turnos,
            "analisis": analisis,
            "lara_ambiguo": bool(notas.get("lara_nota")),
        })


@app.get("/historicos/estudios/2013-municipales", response_class=HTMLResponse)
async def municipales_2013_collection(request: Request):
    """Vista coleccion: 52 exit polls municipales 2013."""
    conn = get_db()
    import json as _json, math as _math

    nac_row = conn.execute(
        "SELECT notas FROM historico_estudios WHERE eleccion_ref='2013-municipales' AND ambito='NACIONAL'"
    ).fetchone()
    if not nac_row:
        conn.close()
        raise HTTPException(404, "Datos de 2013-municipales no encontrados. Ejecute import_2013_municipales.py primero.")
    coleccion_meta = _json.loads(nac_row["notas"] or "{}")

    filas_e = {r["ambito"]: dict(r) for r in conn.execute(
        "SELECT ambito, nombre, pct_gov, pct_opos, pct_otros, num_centros, notas "
        "FROM historico_estudios WHERE eleccion_ref='2013-municipales' AND ambito!='NACIONAL'"
    ).fetchall()}
    filas_o = {r["ambito"]: dict(r) for r in conn.execute(
        "SELECT ambito, pct_gov, pct_opos, pct_otros "
        "FROM historico_oficial WHERE eleccion_ref='2013-municipales'"
    ).fetchall()}
    conn.close()

    ICC_REF = 0.04
    municipios = []
    for slug, fe in filas_e.items():
        fo = filas_o.get(slug, {})
        notas = _json.loads(fe.get("notas") or "{}")
        e_gov = fe.get("pct_gov"); e_opos = fe.get("pct_opos")
        o_gov = fo.get("pct_gov"); o_opos = fo.get("pct_opos")
        n = notas.get("n_respondentes", 0)
        k = fe.get("num_centros", 0)
        deff = round(1 + (n / k - 1) * ICC_REF, 2) if k > 0 else None
        moe = round(1.96 * _math.sqrt((deff or 1) * 0.25 / n) * 100, 2) if n > 0 else None
        win_e = e_gov is not None and e_opos is not None and e_gov > e_opos
        win_o = o_gov is not None and o_opos is not None and o_gov > o_opos
        delta_g = round(e_gov - o_gov, 2) if e_gov is not None and o_gov is not None else None
        delta_o = round(e_opos - o_opos, 2) if e_opos is not None and o_opos is not None else None
        municipios.append({
            "slug": slug,
            "nombre": fe.get("nombre", slug),
            "estado": notas.get("estado", ""),
            "municipio": notas.get("municipio", fe.get("nombre", slug)),
            "e_gov": e_gov, "e_opos": e_opos,
            "o_gov": o_gov, "o_opos": o_opos,
            "n_centros": k, "n_respondentes": n,
            "deff": deff, "moe_pp": moe,
            "win_estudio": win_e, "win_oficial": win_o,
            "ganador_ok": (win_e == win_o) if o_gov is not None else None,
            "delta_gov": delta_g, "delta_opos": delta_o,
            "auditoria": notas.get("auditoria", {}),
            "cand_gov": notas.get("cand_gov_nombre") or "",
            "cand_opos": notas.get("cand_opos_nombre") or "",
        })

    from collections import defaultdict
    por_estado: dict = defaultdict(list)
    for m in municipios:
        por_estado[m["estado"]].append(m)
    estados = [(e, sorted(items, key=lambda x: x["municipio"])) for e, items in sorted(por_estado.items())]
    correctos = sum(1 for m in municipios if m["ganador_ok"] is True)
    con_oficial = sum(1 for m in municipios if m["ganador_ok"] is not None)

    return templates.TemplateResponse(request=request,
        name="municipales_2013.html", context={
            "meta": coleccion_meta,
            "municipios": municipios,
            "estados": estados,
            "total_centros": coleccion_meta.get("n_centros_total", 0),
            "total_resp": coleccion_meta.get("n_respondentes_total", 0),
            "n_estudios": len(municipios),
            "correctos": correctos,
            "con_oficial": con_oficial,
        })


@app.get("/historicos/estudios/2013-municipales/{municipio_slug}", response_class=HTMLResponse)
async def municipales_2013_detalle(request: Request, municipio_slug: str):
    """Detalle individual: tendencia por turno, auditoria y comparacion oficial si existe."""
    conn = get_db()
    import json as _json, math as _math, json as json_mod

    fe = conn.execute(
        "SELECT * FROM historico_estudios WHERE eleccion_ref='2013-municipales' AND ambito=?",
        (municipio_slug,)
    ).fetchone()
    if not fe:
        conn.close()
        raise HTTPException(404, f"Municipio '{municipio_slug}' no encontrado en 2013-municipales")
    fe = dict(fe)
    fo = conn.execute(
        "SELECT * FROM historico_oficial WHERE eleccion_ref='2013-municipales' AND ambito=?",
        (municipio_slug,)
    ).fetchone()
    fo = dict(fo) if fo else {}

    turnos_raw = conn.execute(
        "SELECT turno, hora_label, pct_gov, pct_opos, pct_otros, num_centros "
        "FROM historico_estudios_turnos "
        "WHERE eleccion_ref='2013-municipales' AND ambito=? ORDER BY turno",
        (municipio_slug,)
    ).fetchall()
    conn.close()

    notas = _json.loads(fe.get("notas") or "{}")
    turnos = [dict(r) for r in turnos_raw]
    e_gov = fe.get("pct_gov"); e_opos = fe.get("pct_opos"); e_otros = fe.get("pct_otros")
    o_gov = fo.get("pct_gov"); o_opos = fo.get("pct_opos"); o_otros = fo.get("pct_otros")
    n = notas.get("n_respondentes", 0)
    k = fe.get("num_centros", 0)
    deff = round(1 + (n / k - 1) * 0.04, 2) if k > 0 else None
    moe_srs = round(1.96 * _math.sqrt(0.25 / n) * 100, 2) if n > 0 else None
    moe_adj = round(1.96 * _math.sqrt((deff or 1) * 0.25 / n) * 100, 2) if n > 0 else None

    analisis = {}
    if e_gov is not None and o_gov is not None and o_gov > 0:
        delta_g = round(e_gov - o_gov, 2)
        delta_o = round(e_opos - o_opos, 2) if e_opos is not None and o_opos else None
        e_brecha = round(e_gov - (e_opos or 0), 2)
        o_brecha = round(o_gov - (o_opos or 0), 2)
        error_abs = abs(delta_g)
        magnitud = ("leve" if error_abs <= 1.0 else "moderado" if error_abs <= 2.5
                    else "severo" if error_abs <= 5.0 else "critico")
        win_e = e_gov > (e_opos or 0)
        win_o = o_gov > (o_opos or 0)
        mae_partes = [error_abs]
        if delta_o is not None: mae_partes.append(abs(delta_o))
        if e_otros is not None and o_otros is not None and o_otros > 0:
            mae_partes.append(abs(round(e_otros - o_otros, 2)))
        analisis = {
            "delta_gov": delta_g, "delta_opos": delta_o,
            "e_brecha": e_brecha, "o_brecha": o_brecha,
            "delta_brecha": round(e_brecha - o_brecha, 2),
            "error_abs": error_abs, "magnitud": magnitud,
            "mae_mosteller": round(sum(mae_partes) / len(mae_partes), 2),
            "mae_n_opciones": len(mae_partes),
            "ganador_estudio": "gobierno" if win_e else "oposicion",
            "ganador_oficial": "gobierno" if win_o else "oposicion",
            "acierto_ganador": win_e == win_o,
            "errores_opuestos": (delta_g * delta_o < 0) if delta_o else False,
        }

    return templates.TemplateResponse(request=request,
        name="municipal_2013_detalle.html", context={
            "slug": municipio_slug,
            "nombre": fe.get("nombre", municipio_slug),
            "notas": notas,
            "fe": fe, "fo": fo,
            "e_gov": e_gov, "e_opos": e_opos, "e_otros": e_otros,
            "o_gov": o_gov, "o_opos": o_opos, "o_otros": o_otros,
            "n_respondentes": n, "n_centros": k,
            "deff": deff, "moe_srs": moe_srs, "moe_adj": moe_adj,
            "turnos_json": json_mod.dumps(turnos),
            "turnos": turnos,
            "analisis": analisis,
            "auditoria": notas.get("auditoria", {}),
        })


@app.get("/historicos/estudios/{ref}", response_class=HTMLResponse)
async def historico_estudio_detalle(request: Request, ref: str):
    conn = get_db()
    tiene = (conn.execute("SELECT 1 FROM historico_estudios WHERE eleccion_ref=? LIMIT 1", (ref,)).fetchone()
             or conn.execute("SELECT 1 FROM historico_oficial WHERE eleccion_ref=? LIMIT 1", (ref,)).fetchone())
    if not tiene:
        conn.close()
        raise HTTPException(404, f"No hay datos para '{ref}'")

    pivot = _estudios_pivot(conn, ref)
    nac = next((r for r in pivot if r["ambito"] == "NACIONAL"), {})
    estados_rows = [r for r in pivot if r["ambito"] != "NACIONAL"]
    oficial_solo = nac.get("e_gov") is None and nac.get("o_gov") is not None

    # Metadatos
    meta_e = conn.execute(
        "SELECT nombre_eleccion, fecha_eleccion, notas FROM historico_estudios WHERE eleccion_ref=? AND ambito='NACIONAL'",
        (ref,)
    ).fetchone()
    meta_o = conn.execute(
        "SELECT nombre_eleccion, fecha_eleccion FROM historico_oficial WHERE eleccion_ref=? AND ambito='NACIONAL'",
        (ref,)
    ).fetchone()
    meta = {}
    if meta_e:
        meta.update(dict(meta_e))
    if meta_o:
        for k, v in dict(meta_o).items():
            if not meta.get(k):
                meta[k] = v
    notas = meta.pop("notas", None)
    legislativo = {}
    if notas:
        try:
            parsed_notas = json.loads(notas)
            if parsed_notas.get("tipo") == "asamblea":
                legislativo = parsed_notas
        except Exception:
            legislativo = {}

    # Turnos para el gráfico (presidenciales: ambito=NACIONAL)
    turnos = [dict(r) for r in conn.execute(
        "SELECT turno, hora_label, pct_gov, pct_opos, pct_otros, num_centros "
        "FROM historico_estudios_turnos WHERE eleccion_ref=? AND ambito='NACIONAL' ORDER BY turno",
        (ref,)
    ).fetchall()]

    # Análisis de acertividad
    analisis = {}
    analisis_leg: dict = {}
    if nac.get("e_gov") is not None and nac.get("o_gov") is not None:
        e_g  = nac["e_gov"];  e_o = nac.get("e_opos") or 0
        o_g  = nac["o_gov"];  o_o = nac.get("o_opos") or 0
        delta     = round(e_g - o_g, 2)
        d_opos    = round(e_o - o_o, 2)
        e_brecha  = round(e_g - e_o, 2)
        o_brecha  = round(o_g - o_o, 2)
        d_brecha  = round(e_brecha - o_brecha, 2)
        error_abs = abs(delta)
        magnitud  = ("leve" if error_abs <= 1.0
                     else "moderado" if error_abs <= 2.5
                     else "severo"   if error_abs <= 5.0
                     else "critico")
        ganador_e = "gobierno" if e_g > e_o else "oposición"
        ganador_o = "gobierno" if o_g > o_o else "oposición"
        acierto_ganador = ganador_e == ganador_o
        deltas_estado = [r["delta_gov"] for r in estados_rows if r["delta_gov"] is not None]
        n_est = len(deltas_estado)
        rmse = round((sum(d**2 for d in deltas_estado) / n_est) ** 0.5, 2) if n_est else None
        bias_std = round((sum((d - (sum(deltas_estado)/n_est))**2 for d in deltas_estado) / n_est) ** 0.5, 2) if n_est > 1 else None
        n_neg = sum(1 for d in deltas_estado if d < -0.5)
        pct_neg = round(n_neg / n_est * 100, 1) if n_est else None
        # MAE Mosteller Medida 3: media aritmética de |Δ| por cada opción
        e_ot = nac.get("e_otros"); o_ot = nac.get("o_otros")
        mae_partes = [error_abs, abs(d_opos)]
        if e_ot is not None and o_ot is not None:
            mae_partes.append(abs(round(e_ot - o_ot, 2)))
        mae_mosteller = round(sum(mae_partes) / len(mae_partes), 2)

        analisis = {
            "delta_gov":   delta,    "delta_opos":  d_opos,
            "e_brecha":    e_brecha, "o_brecha":    o_brecha, "delta_brecha": d_brecha,
            "error_abs":   error_abs, "magnitud":   magnitud,
            "mae_mosteller": mae_mosteller, "mae_n_opciones": len(mae_partes),
            "ganador_estudio": ganador_e, "ganador_oficial": ganador_o,
            "acierto_ganador": acierto_ganador,
            "errores_opuestos": (delta * d_opos < 0) if d_opos != 0 else False,
            "rmse_estados": rmse, "bias_std": bias_std,
            "n_estados": n_est, "n_estados_gov_neg": n_neg, "pct_estados_gov_neg": pct_neg,
            "n_centros": int(nac.get("e_centros") or 0),
        }
        # Para legislativo: métricas adicionales sobre escaños + Capa 2 conversión votos→escaños
        if legislativo and legislativo.get("estudio_escanos") and legislativo.get("oficial_escanos"):
            es  = legislativo["estudio_escanos"]
            os_ = legislativo["oficial_escanos"]
            total_esc = os_.get("gov", 0) + os_.get("opos", 0) + os_.get("otros", 0)
            dg_esc = es.get("gov", 0) - os_.get("gov", 0)
            do_esc = es.get("opos", 0) - os_.get("opos", 0)
            err_rel_gov = round(abs(dg_esc) / total_esc * 100, 1) if total_esc else None
            acierto_may = (es.get("gov", 0) >= 83) == (os_.get("gov", 0) >= 83)
            ev = legislativo.get("estudio_voto_lista_pct") or {}
            ov = legislativo.get("oficial_voto_pct") or {}
            d_voto_g = round(ev.get("gov", 0) - ov.get("gov", 0), 2) if ev.get("gov") and ov.get("gov") else None
            d_voto_o = round(ev.get("opos", 0) - ov.get("opos", 0), 2) if ev.get("opos") and ov.get("opos") else None
            mag_esc = ("leve" if abs(dg_esc) <= 5 else "moderado" if abs(dg_esc) <= 15 else "severo")
            # MAE Mosteller M3 para voto lista (Capa 1)
            mae_lista_partes = []
            if d_voto_g is not None: mae_lista_partes.append(abs(d_voto_g))
            if d_voto_o is not None: mae_lista_partes.append(abs(d_voto_o))
            d_voto_ot = round(ev.get("otros", 0) - ov.get("otros", 0), 2) if ev.get("otros") and ov.get("otros") else None
            if d_voto_ot is not None: mae_lista_partes.append(abs(d_voto_ot))
            mae_voto_lista = round(sum(mae_lista_partes) / len(mae_lista_partes), 2) if mae_lista_partes else None
            # Capa 2: descomposición del error de conversión votos → escaños
            # Error proporcional esperado: qué error de escaños produciría puro sesgo de voto
            # Error de algoritmo: el excedente que el sistema electoral añade
            capa2: dict = {}
            if d_voto_g is not None and total_esc > 0:
                err_esc_proporcional = round(abs(d_voto_g) / 100 * total_esc, 1)
                err_esc_algoritmo   = round(abs(dg_esc) - err_esc_proporcional, 1)
                factor_amplif       = round(abs(dg_esc) / err_esc_proporcional, 2) if err_esc_proporcional else None
                capa2 = {
                    "err_esc_real": abs(dg_esc),
                    "err_esc_proporcional": err_esc_proporcional,
                    "err_esc_algoritmo": err_esc_algoritmo,
                    "factor_amplificacion": factor_amplif,
                    "pct_error_por_algoritmo": round(err_esc_algoritmo / abs(dg_esc) * 100, 1) if dg_esc else None,
                }
            analisis_leg = {
                "delta_gov_esc": dg_esc, "delta_opos_esc": do_esc,
                "total_esc": total_esc, "error_rel_gov_pct": err_rel_gov,
                "acierto_mayoria": acierto_may, "magnitud": mag_esc,
                "delta_voto_lista_gov": d_voto_g, "delta_voto_lista_opos": d_voto_o,
                "mae_voto_lista": mae_voto_lista,
                "capa2": capa2,
            }

    # Diagnóstico de sesgo sistémico (solo para tarjetas con estudio)
    sesgo_audit: dict = {}
    if nac.get("e_gov") is not None:
        try:
            from auditor_sesgo import diagnosticar_estudio as _diag
            sesgo_audit = _diag(
                e_gov=nac.get("e_gov"),
                e_opos=nac.get("e_opos"),
                o_gov=nac.get("o_gov"),
                o_opos=nac.get("o_opos"),
                n_centros=int(nac.get("e_centros") or 0),
                estados_rows=estados_rows,
            )
        except Exception:
            sesgo_audit = {"disponible": False, "razon": "error_al_diagnosticar"}

    conn.close()
    import json as _json
    return templates.TemplateResponse(request=request, name="historico_estudio_detalle.html", context={
        "ref": ref, "meta": meta, "nac": nac,
        "estados_rows": estados_rows,
        "turnos_json": _json.dumps(turnos),
        "o_gov_nac": nac.get("o_gov"),
        "o_opos_nac": nac.get("o_opos"),
        "oficial_solo": oficial_solo,
        "analisis": analisis,
        "analisis_leg": analisis_leg,
        "legislativo": legislativo,
        "sesgo_audit": sesgo_audit,
    })


@app.get("/historicos/{ref}", response_class=HTMLResponse)
async def historico_detalle(request: Request, ref: str):
    conn = get_db()
    if not conn.execute(
        "SELECT 1 FROM resultados_historicos WHERE eleccion_ref=? LIMIT 1", (ref,)
    ).fetchone():
        conn.close()
        raise HTTPException(404, detail=f"No hay datos para '{ref}'")

    resumen = conn.execute("""
        SELECT
            COUNT(*)                                                            AS num_centros,
            SUM(votos_validos)                                                  AS total_votos,
            ROUND(SUM(votos_gobierno)*100.0  / NULLIF(SUM(votos_validos),0),1) AS pct_gov,
            ROUND(SUM(votos_oposicion)*100.0 / NULLIF(SUM(votos_validos),0),1) AS pct_opos
        FROM resultados_historicos
        WHERE eleccion_ref = ?
    """, (ref,)).fetchone()

    estados = conn.execute("""
        SELECT
            e.nombre                                                            AS estado,
            COUNT(*)                                                            AS num_centros,
            SUM(rh.votos_validos)                                               AS total_votos,
            ROUND(SUM(rh.votos_gobierno)*100.0  / NULLIF(SUM(rh.votos_validos),0),1) AS pct_gov,
            ROUND(SUM(rh.votos_oposicion)*100.0 / NULLIF(SUM(rh.votos_validos),0),1) AS pct_opos
        FROM resultados_historicos rh
        JOIN centros c ON rh.codigo_centro = c.codigo_cne
        JOIN estados e ON c.id_estado = e.id
        WHERE rh.eleccion_ref = ?
        GROUP BY e.id
        ORDER BY e.nombre
    """, (ref,)).fetchall()

    top_centros = conn.execute("""
        SELECT
            rh.codigo_centro,
            c.nombre        AS centro,
            e.nombre        AS estado,
            rh.votos_validos,
            rh.pct_gobierno,
            rh.pct_oposicion
        FROM resultados_historicos rh
        JOIN centros c ON rh.codigo_centro = c.codigo_cne
        JOIN estados e ON c.id_estado = e.id
        WHERE rh.eleccion_ref = ?
        ORDER BY rh.votos_validos DESC
        LIMIT 20
    """, (ref,)).fetchall()

    safe_ref = re.sub(r"[^a-zA-Z0-9_-]", "_", ref)
    mapa_existe = (VIZ_DIR / f"hist_{safe_ref}.html").exists()
    conn.close()

    return templates.TemplateResponse(request=request, name="historico_detalle.html", context={
        "ref": ref, "resumen": resumen, "estados": estados,
        "top_centros": top_centros, "mapa_existe": mapa_existe,
    })


@app.get("/historicos/{ref}/mapa")
async def historico_mapa(ref: str):
    conn = get_db()
    if not conn.execute(
        "SELECT 1 FROM resultados_historicos WHERE eleccion_ref=? LIMIT 1", (ref,)
    ).fetchone():
        conn.close()
        raise HTTPException(404)

    datos_ventaja = _datos_ventaja_historico_ref(conn, ref)
    conn.close()

    if not datos_ventaja:
        return RedirectResponse(f"/historicos/{ref}", status_code=303)

    safe_ref = re.sub(r"[^a-zA-Z0-9_-]", "_", ref)
    ruta = str(VIZ_DIR / f"hist_{safe_ref}.html")

    sys.path.insert(0, str(BASE_DIR))
    import generador_heatmap
    import importlib
    importlib.reload(generador_heatmap)
    generador_heatmap.generar_heatmap(datos_ventaja, nivel="estado", ruta_salida=ruta, titulo=ref)

    return RedirectResponse(f"/static/viz/hist_{safe_ref}.html", status_code=303)


# ── Tests / demo data ─────────────────────────────────────────────────────────

@app.post("/test/demo")
async def test_demo():
    """Carga el dataset completo (todos los turnos) con datos CNE 2024."""
    db = get_db()
    try:
        eleccion = db.execute("SELECT * FROM elecciones WHERE activa=1").fetchone()
        if not eleccion:
            return JSONResponse({"ok": False, "mensaje": "No hay elección activa."}, status_code=400)
        cands   = _asegurar_candidatos(db, eleccion["id"])
        centros = db.execute(
            "SELECT m.id AS id_muestra, m.codigo_centro, c.num_electores, c.lat, c.lon "
            "FROM muestra m JOIN centros c ON c.codigo_cne=m.codigo_centro "
            "WHERE m.id_eleccion=? AND m.activo=1 AND c.activo=1", (eleccion["id"],)
        ).fetchall()
        pct = _pct_2024()
        db.execute("DELETE FROM votos"); db.execute("DELETE FROM sms_raw"); db.commit()
        _encuestadores_demo(db, centros, eleccion["id"])
        turnos = _turnos(eleccion)
        _insertar_votos(db, eleccion, centros, cands, pct, turnos)
        db.commit()
        total = db.execute("SELECT COUNT(*) FROM votos").fetchone()[0]
        return JSONResponse({"ok": True, "mensaje": f"Dataset completo cargado: {total:,} opiniones en {len(turnos)} turnos."})
    except Exception as e:
        db.rollback()
        return JSONResponse({"ok": False, "mensaje": str(e)}, status_code=500)
    finally:
        db.close()


@app.post("/test/entrada")
async def test_entrada():
    """Carga solo los primeros 6 turnos (~2 horas) para mostrar el estado de muestra parcial."""
    db = get_db()
    try:
        eleccion = db.execute("SELECT * FROM elecciones WHERE activa=1").fetchone()
        if not eleccion:
            return JSONResponse({"ok": False, "mensaje": "No hay elección activa."}, status_code=400)
        cands   = _asegurar_candidatos(db, eleccion["id"])
        centros = db.execute(
            "SELECT m.id AS id_muestra, m.codigo_centro, c.num_electores, c.lat, c.lon "
            "FROM muestra m JOIN centros c ON c.codigo_cne=m.codigo_centro "
            "WHERE m.id_eleccion=? AND m.activo=1 AND c.activo=1", (eleccion["id"],)
        ).fetchall()
        pct = _pct_2024()
        db.execute("DELETE FROM votos"); db.execute("DELETE FROM sms_raw"); db.commit()
        _encuestadores_demo(db, centros, eleccion["id"])
        # Solo la mitad de centros reportan (simula entrada parcial de datos)
        centros_activos = centros[: len(centros) // 2]
        turnos = _turnos(eleccion)[:6]
        _insertar_votos(db, eleccion, centros_activos, cands, pct, turnos)
        db.commit()
        total = db.execute("SELECT COUNT(*) FROM votos").fetchone()[0]
        return JSONResponse({"ok": True, "mensaje": f"Entrada parcial cargada: {total:,} opiniones - {len(centros_activos)}/{len(centros)} centros, primeros 6 turnos."})
    except Exception as e:
        db.rollback()
        return JSONResponse({"ok": False, "mensaje": str(e)}, status_code=500)
    finally:
        db.close()


@app.post("/test/reset")
async def test_reset():
    """Elimina todas las opiniones y sms_raw de la elección activa."""
    db = get_db()
    try:
        db.execute("DELETE FROM votos")
        db.execute("DELETE FROM sms_raw")
        db.commit()
        return JSONResponse({"ok": True, "mensaje": "Datos de prueba eliminados."})
    except Exception as e:
        db.rollback()
        return JSONResponse({"ok": False, "mensaje": str(e)}, status_code=500)
    finally:
        db.close()
